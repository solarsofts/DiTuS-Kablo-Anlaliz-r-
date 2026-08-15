from __future__ import annotations

from copy import deepcopy
from dataclasses import asdict, dataclass, field, replace
from datetime import date, datetime
from hashlib import sha256
from html import escape
from pathlib import Path
import json
import math
import re
from typing import Any, Iterable, Mapping, Sequence

from ucd.calculations.project_application import (
    calculate_project_voltage_drop,
    evaluate_application_iteration_gates,
)
from ucd.calculations.source_audit import audit_project_sources
from ucd.calculations.procurement import build_procurement_package
from ucd.calculations.bonding_accessories import resolve_bonding_accessory_plan
from ucd import __version__
from ucd.models.project import ProjectData


REFERENCE = (
    f"DiTuS v{__version__} project, calculation and procurement-summary reporting engine. Reports reproduce the "
    "project state and calculation results supplied to the engine; they do not turn "
    "conditional or incomplete calculations into final design approval."
)

REPORT_CALCULATION = "CALCULATION_REPORT"
REPORT_DESIGN = "DESIGN_REPORT"
REPORT_SUMMARY = "TECHNICAL_SUMMARY"
REPORT_FULL = "FULL_DESIGN_DOSSIER"

MODULE_PROJECT = "PROJECT_SUMMARY"
MODULE_DESIGN_BASIS = "DESIGN_BASIS"
MODULE_SOURCE_AUDIT = "SOURCE_AUDIT"
MODULE_CABLE = "CABLE_CONSTRUCTION"
MODULE_ROUTE = "ROUTE_AND_REGIONS"
MODULE_VOLTAGE_DROP = "VOLTAGE_DROP"
MODULE_IEC60287 = "IEC60287"
MODULE_NODAL = "NODAL_THERMAL"
MODULE_TRANSIENT = "TRANSIENT_THERMAL"
MODULE_BONDING = "BONDING"
MODULE_FAULT = "FAULT_EPR"
MODULE_SVL = "SVL"
MODULE_PROCUREMENT = "PROCUREMENT_SUMMARY"
MODULE_TRACE = "CALCULATION_TRACE"
MODULE_WARNINGS = "WARNINGS_LIMITATIONS"

MODULE_LABELS: dict[str, str] = {
    MODULE_PROJECT: "Proje özeti",
    MODULE_DESIGN_BASIS: "Tasarım esası ve yük senaryoları",
    MODULE_SOURCE_AUDIT: "Kaynak veri ve çelişki denetimi",
    MODULE_CABLE: "Kablo konstrüksiyonu ve veri kaynağı",
    MODULE_ROUTE: "Güzergâh ve termal bölgeler",
    MODULE_VOLTAGE_DROP: "Gerilim düşümü ön hesabı",
    MODULE_IEC60287: "IEC 60287 sürekli akım taşıma kapasitesi",
    MODULE_NODAL: "2D nodal kararlı termal analiz",
    MODULE_TRANSIENT: "IEC 60853 geçici/çevrimsel analiz",
    MODULE_BONDING: "Metalik kılıf ve bonding",
    MODULE_FAULT: "Arıza, EPR ve metalik kılıf dayanımı",
    MODULE_SVL: "SVL ve yalıtım koordinasyonu",
    MODULE_PROCUREMENT: "BOQ/BOM ve tedarik özeti",
    MODULE_TRACE: "Hesap izi ve yöntem referansları",
    MODULE_WARNINGS: "Uyarılar, sınırlamalar ve açık işler",
}

ALL_MODULES: tuple[str, ...] = tuple(MODULE_LABELS)

REPORT_TEMPLATES: dict[str, tuple[str, ...]] = {
    REPORT_CALCULATION: (
        MODULE_PROJECT,
        MODULE_DESIGN_BASIS,
        MODULE_SOURCE_AUDIT,
        MODULE_CABLE,
        MODULE_ROUTE,
        MODULE_VOLTAGE_DROP,
        MODULE_IEC60287,
        MODULE_NODAL,
        MODULE_TRANSIENT,
        MODULE_BONDING,
        MODULE_FAULT,
        MODULE_SVL,
        MODULE_PROCUREMENT,
        MODULE_TRACE,
        MODULE_WARNINGS,
    ),
    REPORT_DESIGN: (
        MODULE_PROJECT,
        MODULE_DESIGN_BASIS,
        MODULE_SOURCE_AUDIT,
        MODULE_CABLE,
        MODULE_ROUTE,
        MODULE_VOLTAGE_DROP,
        MODULE_IEC60287,
        MODULE_NODAL,
        MODULE_BONDING,
        MODULE_FAULT,
        MODULE_SVL,
        MODULE_PROCUREMENT,
        MODULE_WARNINGS,
    ),
    REPORT_SUMMARY: (
        MODULE_PROJECT,
        MODULE_DESIGN_BASIS,
        MODULE_CABLE,
        MODULE_ROUTE,
        MODULE_VOLTAGE_DROP,
        MODULE_WARNINGS,
    ),
    REPORT_FULL: ALL_MODULES,
}

REPORT_TYPE_LABELS: dict[str, str] = {
    REPORT_CALCULATION: "Hesap Raporu",
    REPORT_DESIGN: "Proje / Tasarım Raporu",
    REPORT_SUMMARY: "Kısa Teknik Özet",
    REPORT_FULL: "Tam Tasarım Dosyası",
}


@dataclass(frozen=True)
class ReportMetadata:
    report_type: str = REPORT_DESIGN
    title: str = ""
    document_no: str = ""
    revision: str = "00"
    issue_date: str = field(default_factory=lambda: date.today().isoformat())
    client: str = ""
    contractor: str = ""
    prepared_by: str = ""
    checked_by: str = ""
    approval_status: str = "TASLAK"
    confidentiality: str = ""


@dataclass(frozen=True)
class ReportConfiguration:
    metadata: ReportMetadata = field(default_factory=ReportMetadata)
    selected_modules: tuple[str, ...] = field(default_factory=tuple)
    include_detailed_trace: bool = False
    include_empty_selected_modules: bool = True
    output_formats: tuple[str, ...] = ("html", "markdown", "json", "docx", "pdf")

    def resolved_modules(self) -> tuple[str, ...]:
        requested = self.selected_modules or REPORT_TEMPLATES.get(
            self.metadata.report_type, REPORT_TEMPLATES[REPORT_DESIGN]
        )
        ordered: list[str] = []
        for module in ALL_MODULES:
            if module in requested and module not in ordered:
                ordered.append(module)
        if MODULE_WARNINGS not in ordered:
            ordered.append(MODULE_WARNINGS)
        return tuple(ordered)


@dataclass(frozen=True)
class ReportTable:
    title: str
    headers: tuple[str, ...]
    rows: tuple[tuple[str, ...], ...]
    notes: tuple[str, ...] = ()


@dataclass(frozen=True)
class ReportSection:
    section_id: str
    title: str
    status: str = "AVAILABLE"
    paragraphs: tuple[str, ...] = ()
    tables: tuple[ReportTable, ...] = ()
    warnings: tuple[str, ...] = ()
    trace: tuple[str, ...] = ()


@dataclass(frozen=True)
class CalculationResultsBundle:
    iec_results: tuple[Any, ...] = ()
    production_electrothermal_result: Any | None = None
    production_bonding_result: Any | None = None
    nodal_thermal_result: Any | None = None
    transient_thermal_result: Any | None = None
    bonding_result: Any | None = None
    fault_result: Any | None = None
    svl_result: Any | None = None


@dataclass(frozen=True)
class ProjectReport:
    reference: str
    generated_at: str
    project_name: str
    project_code: str
    project_signature_sha256: str
    metadata: ReportMetadata
    selected_modules: tuple[str, ...]
    sections: tuple[ReportSection, ...]
    mandatory_warnings: tuple[str, ...]
    report_status: str
    trace: tuple[str, ...]

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


class ReportGenerationError(ValueError):
    pass


def default_report_configuration(report_type: str = REPORT_DESIGN) -> ReportConfiguration:
    return ReportConfiguration(metadata=ReportMetadata(report_type=report_type))


def _stable_project_signature(project: ProjectData) -> str:
    payload = asdict(project)
    payload.pop("modified_at", None)
    raw = json.dumps(payload, ensure_ascii=False, sort_keys=True, separators=(",", ":"), default=str)
    return sha256(raw.encode("utf-8")).hexdigest()


def _fmt(value: Any, decimals: int = 3, suffix: str = "") -> str:
    if value is None:
        return "—"
    if isinstance(value, bool):
        return "Evet" if value else "Hayır"
    if isinstance(value, int):
        return f"{value}{suffix}"
    if isinstance(value, float):
        if not math.isfinite(value):
            return "—"
        return f"{value:.{decimals}f}{suffix}"
    text = str(value).strip()
    return text if text else "—"


def _status_text(value: Any) -> str:
    text = str(value or "").strip()
    return text or "BELİRTİLMEMİŞ"


def _rows_from_mapping(mapping: Mapping[str, Any], labels: Mapping[str, str]) -> tuple[tuple[str, str], ...]:
    return tuple((labels.get(key, key), _fmt(mapping.get(key))) for key in labels)


def _section_project(project: ProjectData) -> ReportSection:
    metadata_rows = (
        ("Proje adı", project.project_name),
        ("Proje kodu", project.project_code),
        ("Açıklama", project.description or "—"),
        ("Standart profili", project.standards_profile),
        ("Proje şema sürümü", project.schema_version),
        ("Oluşturma zamanı", project.created_at),
        ("Son değişiklik", project.modified_at),
        ("Kapsam", project.source_audit.scope or "UNDERGROUND_ONLY"),
    )
    return ReportSection(
        MODULE_PROJECT,
        MODULE_LABELS[MODULE_PROJECT],
        paragraphs=(
            "Bu rapor, rapor üretildiği anda DiTuS proje modelinde bulunan verileri ve bu oturumda "
            "rapor motoruna aktarılan hesap sonuçlarını esas alır.",
        ),
        tables=(ReportTable("Proje kimliği", ("Alan", "Değer"), metadata_rows),),
    )


def _section_design_basis(project: ProjectData) -> ReportSection:
    d = project.design_basis
    rows = (
        ("Nominal sistem gerilimi", _fmt(d.system_voltage_kv, 3, " kV")),
        ("Frekans", _fmt(d.frequency_hz, 2, " Hz")),
        ("Devre sayısı", _fmt(d.circuit_count, 0)),
        ("Aktif devre sayısı", _fmt(d.active_circuit_count, 0)),
        ("N-1 etkin", _fmt(d.n_minus_one_enabled)),
        ("Topraklama tipi", d.grounding_type),
        ("Yük giriş modu", d.load_input_mode),
        ("Aktif güç", _fmt(d.active_power_mw, 3, " MW")),
        ("Görünür güç", _fmt(d.apparent_power_mva, 3, " MVA")),
        ("Güç faktörü", _fmt(d.power_factor, 4)),
        ("Normal toplam akım", _fmt(d.normal_total_current_a, 3, " A")),
        ("Normal akım / aktif devre", _fmt(d.normal_current_per_active_circuit_a, 3, " A")),
        ("N-1 akımı / devre", _fmt(d.n1_current_per_circuit_a, 3, " A")),
        ("Tasarım akımı / devre", _fmt(d.design_current_per_circuit_a, 3, " A")),
        ("Tasarım marjı", _fmt(d.design_margin_percent, 2, " %")),
        ("Toplam yeraltı güzergâhı", _fmt(sum(r.length_m for r in project.route_sections), 3, " m")),
        ("Başlangıç kurulum profili", d.installation_profile),
        ("Toprak ısıl özdirenci", _fmt(d.soil_thermal_resistivity_km_w, 4, " K.m/W")),
        ("Toprak verisi kaynağı", d.soil_thermal_value_source),
    )
    warnings = []
    if d.design_current_per_circuit_a <= 0:
        warnings.append("Tasarım akımı pozitif bir değer olarak tanımlanmamıştır.")
    if d.soil_thermal_value_source in {"PRELIMINARY_ASSUMPTION", "UNKNOWN", ""}:
        warnings.append("Toprak ısıl özdirenci ölçülmüş/as-built veri değildir.")
    return ReportSection(
        MODULE_DESIGN_BASIS,
        MODULE_LABELS[MODULE_DESIGN_BASIS],
        tables=(ReportTable("Tasarım girdileri", ("Parametre", "Değer"), rows),),
        warnings=tuple(warnings),
    )


def _section_source_audit(project: ProjectData) -> ReportSection:
    report = audit_project_sources(project)
    issue_rows: list[tuple[str, ...]] = []
    for issue in report.issues:
        issue_rows.append((
            _status_text(getattr(issue, "severity", "")),
            _status_text(getattr(issue, "title", "")),
            _status_text(getattr(issue, "notes", "") or ", ".join(getattr(issue, "source_references", ()))),
        ))
    missing_rows = tuple((str(item),) for item in project.source_audit.missing_required_data)
    warnings: list[str] = []
    if report.status != "PASS" or report.issue_count or report.missing_data_count:
        warnings.extend([
            f"Kaynak veri denetimi durumu: {report.status}",
            f"Açık çelişki/uyarı sayısı: {report.issue_count}",
            f"Nihai tasarım için eksik veri sayısı: {report.missing_data_count}",
        ])
    tables = [
        ReportTable(
            "Kaynak denetimi özeti",
            ("Alan", "Değer"),
            (
                ("Kaynak adı", project.source_audit.source_name or "—"),
                ("Kaynak dosya", project.source_audit.source_file or "—"),
                ("Kapsam", project.source_audit.scope or "—"),
                ("Kapsam dışı alanlar", ", ".join(project.source_audit.excluded_scopes) or "—"),
                ("Denetim durumu", report.status),
            ),
        )
    ]
    if issue_rows:
        tables.append(ReportTable("Kaynak çelişkileri ve denetim bulguları", ("Önem", "Başlık", "Açıklama"), tuple(issue_rows)))
    if missing_rows:
        tables.append(ReportTable("Nihai tasarım için eksik veriler", ("Eksik veri",), missing_rows))
    return ReportSection(
        MODULE_SOURCE_AUDIT,
        MODULE_LABELS[MODULE_SOURCE_AUDIT],
        status=report.status,
        tables=tuple(tables),
        warnings=tuple(warnings),
    )


def _section_cable(project: ProjectData) -> ReportSection:
    c = project.cable
    summary_rows = (
        ("Kablo adı", c.name or c.model or "—"),
        ("Üretici", c.manufacturer or "—"),
        ("Seri / model", " / ".join(part for part in (c.series, c.model) if part) or "—"),
        ("Gerilim sınıfı", c.voltage_class or "—"),
        ("İletken", f"{c.conductor_material} {_fmt(c.conductor_area_mm2, 1)} mm²"),
        ("Paralel kablo / faz", _fmt(c.parallel_cables_per_phase, 0)),
        ("Metalik ekran/kılıf kesiti", _fmt(c.sheath_cross_section_mm2, 3, " mm²")),
        ("Kablo dış çapı", _fmt(c.overall_diameter_mm, 3, " mm")),
        ("Rdc20", _fmt(c.dc_resistance_20_ohm_km, 8, " Ω/km")),
        ("Kapasitans", _fmt(c.capacitance_uf_km, 6, " µF/km")),
        ("Azami iletken sıcaklığı", _fmt(c.max_temperature_c, 2, " °C")),
        ("Veri durumu", c.data_status),
        ("Katalog kayıt kimliği", c.catalog_record_id or "—"),
        ("Proje kablosu kayıt kimliği", c.snapshot_id or "—"),
        ("Proje kablosu veri imzası", c.snapshot_hash or "—"),
    )
    layer_rows = []
    for layer in c.layers:
        layer_rows.append((
            layer.layer_id,
            layer.name,
            layer.layer_type,
            _fmt(layer.inner_diameter_mm, 3),
            _fmt(layer.outer_diameter_mm, 3),
            layer.material or "—",
            _fmt(layer.thermal_resistivity_km_w, 4),
            _fmt(layer.conductor_area_mm2, 3),
            layer.source_id or "—",
        ))
    source_rows = []
    for source in c.parameter_sources:
        source_rows.append((
            source.source_id,
            source.source_type,
            source.document_title or "—",
            source.document_revision or "—",
            source.page_reference or "—",
            "VERIFIED" if source.verified else "UNVERIFIED",
        ))
    warnings = []
    if not c.snapshot_hash:
        warnings.append("Kablo, proje içinde sabit bir proje kaydı olarak imzalanmamıştır.")
    if c.data_status not in {"VALIDATED", "AS_BUILT", "TESTED"}:
        warnings.append(f"Kablo veri durumu nihai doğrulanmış değildir: {c.data_status}")
    tables: list[ReportTable] = [ReportTable("Kablo özeti", ("Parametre", "Değer"), summary_rows)]
    if layer_rows:
        tables.append(ReportTable(
            "Parametrik kablo katmanları",
            ("ID", "Katman", "Tip", "İç çap [mm]", "Dış çap [mm]", "Malzeme", "ρth [K.m/W]", "Alan [mm²]", "Kaynak"),
            tuple(layer_rows),
        ))
    if source_rows:
        tables.append(ReportTable(
            "Kablo veri kaynakları",
            ("Kaynak ID", "Tür", "Doküman", "Revizyon", "Sayfa", "Doğrulama"),
            tuple(source_rows),
        ))
    return ReportSection(MODULE_CABLE, MODULE_LABELS[MODULE_CABLE], tables=tuple(tables), warnings=tuple(warnings))


def _section_route(project: ProjectData) -> ReportSection:
    route_rows = []
    chainage = 0.0
    for index, section in enumerate(project.route_sections, start=1):
        start = section.start_chainage_m if section.end_chainage_m > section.start_chainage_m else chainage
        end = section.end_chainage_m if section.end_chainage_m > start else start + section.length_m
        route_rows.append((
            str(index),
            section.name,
            section.section_type,
            _fmt(start, 3),
            _fmt(end, 3),
            _fmt(section.length_m, 3),
            _fmt(section.burial_depth_m, 3),
            _fmt(section.phase_spacing_m, 3),
            _fmt(section.soil_thermal_resistivity_km_w, 4),
            section.thermal_data_state,
            section.notes or "—",
        ))
        chainage = end
    region_rows = []
    for region in project.thermal_design.regions:
        region_rows.append((
            region.region_id,
            region.name,
            _fmt(region.start_m, 3),
            _fmt(region.end_m, 3),
            region.template_id,
            region.data_state,
            region.source_reference or "—",
        ))
    tables = [ReportTable(
        "Yeraltı güzergâh bölümleri",
        ("No", "Bölüm", "Tip", "Başlangıç [m]", "Bitiş [m]", "Uzunluk [m]", "Derinlik [m]", "Faz aralığı [m]", "Toprak ρth", "Veri durumu", "Not"),
        tuple(route_rows),
    )]
    if region_rows:
        tables.append(ReportTable(
            "Termal bölgeler",
            ("Bölge ID", "Ad", "Başlangıç [m]", "Bitiş [m]", "Kesit şablonu", "Durum", "Kaynak"),
            tuple(region_rows),
        ))
    warnings = []
    if not project.route_sections:
        warnings.append("Yeraltı güzergâh bölümü tanımlanmamıştır.")
    return ReportSection(MODULE_ROUTE, MODULE_LABELS[MODULE_ROUTE], tables=tuple(tables), warnings=tuple(warnings))


def _section_voltage_drop(project: ProjectData) -> ReportSection:
    try:
        result = calculate_project_voltage_drop(deepcopy(project))
    except Exception as exc:  # Calculation may be intentionally blocked by incomplete project data.
        return ReportSection(
            MODULE_VOLTAGE_DROP,
            MODULE_LABELS[MODULE_VOLTAGE_DROP],
            status="NOT_AVAILABLE",
            paragraphs=("Gerilim düşümü ön hesabı üretilemedi.",),
            warnings=(str(exc),),
        )
    rows = (
        ("Toplam güzergâh", _fmt(result.length_m, 3, " m")),
        ("Kablo başına akım", _fmt(result.current_per_cable_a, 3, " A")),
        ("R @ işletme", _fmt(result.resistance_ohm_km, 8, " Ω/km")),
        ("X", _fmt(result.reactance_ohm_km, 8, " Ω/km")),
        ("Gerilim düşümü", _fmt(result.voltage_drop_v, 5, " V")),
        ("Gerilim düşümü", _fmt(result.voltage_drop_percent, 6, " %")),
        ("Kaynak", result.source),
    )
    return ReportSection(
        MODULE_VOLTAGE_DROP,
        MODULE_LABELS[MODULE_VOLTAGE_DROP],
        paragraphs=("Bu bölüm katalog/proje R-L verileriyle yapılan dengeli üç faz ön hesabıdır; termal rating veya nihai uygunluk kararı değildir.",),
        tables=(ReportTable("Gerilim düşümü", ("Parametre", "Değer"), rows),),
        warnings=(),
        trace=(f"Kaynak: {result.source}",),
    )


def _not_run(module: str) -> ReportSection:
    return ReportSection(
        module,
        MODULE_LABELS[module],
        status="NOT_RUN",
        paragraphs=("Bu hesap modülünün sonucu rapor üretim oturumuna aktarılmamıştır.",),
        warnings=("Modül seçilmiş olmasına rağmen hesap sonucu bulunmuyor; rapor nihai hesap kanıtı olarak kullanılamaz.",),
    )


def _section_iec(results: Sequence[Any], production: Any | None = None) -> ReportSection:
    if not results and production is None:
        return _not_run(MODULE_IEC60287)
    rows = []
    warnings = []
    trace = []
    for item in results:
        rows.append((
            _fmt(getattr(item, "section_name", "—")),
            _fmt(getattr(item, "design_current_a", None), 3),
            _fmt(getattr(item, "ampacity_a", None), 3),
            _fmt(getattr(item, "margin_a", None), 3),
            _fmt(getattr(item, "conductor_temperature_at_design_c", None), 3),
            _fmt(getattr(item, "temperature_limit_c", None), 3),
            _fmt(getattr(item, "total_loss_at_design_w_m", None), 4),
            _fmt(getattr(item, "sheath_loss_factor", None), 6),
            _fmt(getattr(item, "armour_loss_factor", None), 6),
            _status_text(getattr(item, "status", "")),
        ))
        warnings.extend(str(x) for x in getattr(item, "notes", ()) if str(x).strip())
        if hasattr(item, "trace_lines"):
            trace.extend(str(x) for x in item.trace_lines())
    tables = []
    if rows:
        tables.append(ReportTable(
            "IEC 60287 bölüm sonuçları",
            ("Bölüm", "I tasarım [A]", "Ampacity [A]", "Marj [A]", "Tcond [°C]", "Limit [°C]", "Toplam kayıp [W/m]", "Legacy λ1 girdisi", "λ2", "Durum"),
            tuple(rows),
        ))
    section_status = "AVAILABLE"
    if production is not None:
        scenario_rows = []
        for scenario_result in getattr(production, "scenarios", ()):
            scenario = getattr(scenario_result, "scenario", None)
            circuit_summary = ", ".join(
                f"{getattr(item, 'circuit_id', '?')}={'OFF' if not getattr(item, 'energized', False) else f'{getattr(item, 'phase_current_a', 0.0):.3f} A'}"
                for item in getattr(scenario, "circuit_states", ())
            )
            aliases = ", ".join(getattr(scenario, "equivalent_scenario_ids", ()) or ())
            scenario_rows.append((
                _status_text(getattr(scenario, "scenario_id", "")),
                _status_text(getattr(scenario, "scenario_name", "")),
                circuit_summary or "—",
                aliases or "—",
                _status_text(getattr(scenario_result, "thermal_method", "ANALYTIC")),
                ", ".join(getattr(scenario_result, "dryout_material_ids", ()) or ()) or "—",
                _status_text(getattr(scenario_result, "completion_status", "")),
                _status_text(getattr(scenario_result, "suitability_status", "")),
                _fmt(getattr(scenario_result, "maximum_conductor_temperature_c", None), 3),
                _status_text(getattr(scenario_result, "critical_region_id", "")),
                _status_text(getattr(scenario_result, "critical_physical_cable_id", "")),
                _fmt(getattr(scenario_result, "network_sheath_loss_ratio", None), 7),
                _fmt(getattr(scenario_result, "lambda1_eddy", None), 7),
                _fmt(getattr(scenario_result, "lambda1_rating", getattr(scenario_result, "global_lambda1", None)), 7),
                _status_text(getattr(scenario_result, "sheath_loss_authority", "UNKNOWN")),
                ", ".join(getattr(scenario_result, "sheath_loss_sources", ()) or ()) or "—",
                ", ".join(getattr(scenario_result, "sheath_loss_reason_codes", ()) or ()) or "—",
                _status_text(getattr(scenario_result, "loss_vector_fingerprint", "")),
            ))
            trace.extend(str(x) for x in getattr(scenario_result, "trace", ()))
            if getattr(scenario_result, "completion_status", "") != "COMPLETE":
                warnings.append(
                    f"{getattr(scenario, 'scenario_id', 'Senaryo')}: üretim elektro-termal çalışma noktası tamamlanmadı."
                )
            if getattr(scenario_result, "sheath_loss_authority", "FULL") != "FULL":
                warnings.append(
                    f"{getattr(scenario, 'scenario_id', 'Senaryo')}: kılıf kaybı fiziği tamamlanmamış; IEC ampacity üretim otoritesi yok. "
                    f"Neden={','.join(getattr(scenario_result, 'sheath_loss_reason_codes', ()) or ()) or 'UNKNOWN'}."
                )
            if getattr(scenario_result, "suitability_status", "") == "UYGUN_DEGIL":
                warnings.append(
                    f"{getattr(scenario, 'scenario_id', 'Senaryo')}: çözülen fiziksel çalışma noktası sıcaklık sınırını aşıyor."
                )
        if scenario_rows:
            tables.append(ReportTable(
                "Senaryo bazlı üretim elektro-termal çalışma noktaları",
                ("Senaryo", "Ad", "Devre akım/enerji durumu", "Kapsanan kimlikler", "Termal yöntem", "Kuruma malzemeleri", "Tamamlanma", "Uygunluk", "Tcond max [°C]", "Kritik bölge", "Kritik kablo", "λ1′ network", "λ1″ eddy", "λ1 rating", "Sheath-loss authority", "λ1″ kaynağı", "Neden kodları", "Kayıp vektörü"),
                tuple(scenario_rows),
                notes=(
                    "λ1′ network, global primitive ağın boyuna metalik kılıf I²R kayıp oranıdır; IEC kapalı-form λ1′ ile cebirsel özdeşlik iddiası yoktur.",
                    "λ1 rating = network boyuna kılıf kaybı + uygulanabilir IEC/izlenebilir harici λ1″ eddy-current bileşenidir.",
                    "Devre dışı kablolar fiziksel geometriden silinmez, ancak iletken/kılıf/dielektrik kayıpları sıfırlanır.",
                    "Kritik-izoterm kuruma verisi etkinse üretim termal yöntemi AUTO tarafından nodal çözüme yükseltilir.",
                ),
            ))
        if any(getattr(item, "completion_status", "") == "FAILED" for item in getattr(production, "scenarios", ())):
            section_status = "FAIL"
        elif warnings:
            section_status = "CONDITIONAL"
    return ReportSection(
        MODULE_IEC60287,
        MODULE_LABELS[MODULE_IEC60287],
        status=section_status,
        tables=tuple(tables),
        warnings=tuple(dict.fromkeys(warnings)),
        trace=tuple(trace),
    )


def _section_nodal(result: Any | None) -> ReportSection:
    if result is None:
        return _not_run(MODULE_NODAL)
    scenarios = getattr(result, "scenarios", ())
    rows = []
    region_rows = []
    warnings = []
    trace = []
    for scenario in scenarios:
        rows.append((
            _status_text(getattr(scenario, "scenario_id", "")),
            _status_text(getattr(scenario, "scenario_name", "")),
            _fmt(getattr(scenario, "current_per_cable_a", None), 3),
            _fmt(getattr(scenario, "route_ampacity_per_cable_a", None), 3),
            _status_text(getattr(scenario, "critical_region_name", "")),
            _fmt(getattr(scenario, "maximum_conductor_temperature_c", None), 3),
            _status_text(getattr(scenario, "status", "")),
        ))
        trace.extend(str(x) for x in getattr(scenario, "trace", ()))
        for region in getattr(scenario, "regions", ()):
            region_rows.append((
                _status_text(getattr(scenario, "scenario_id", "")),
                _status_text(getattr(region, "region_id", "")),
                _status_text(getattr(region, "region_name", "")),
                _fmt(getattr(region, "start_m", None), 3),
                _fmt(getattr(region, "end_m", None), 3),
                _fmt(getattr(region, "design_current_per_cable_a", None), 3),
                _fmt(getattr(region, "ampacity_per_cable_a", None), 3),
                _fmt(getattr(region, "maximum_conductor_temperature_c", None), 3),
                _fmt(getattr(region, "energy_balance_error_percent", None), 5),
                f"{_fmt(getattr(region, 'mesh_nx', None), 0)}×{_fmt(getattr(region, 'mesh_ny', None), 0)}",
                (
                    f"{getattr(region, 'dryout_cell_count', 0)}/{getattr(region, 'dryout_eligible_cell_count', 0)} "
                    f"(%{100.0 * float(getattr(region, 'dryout_fraction', 0.0)):.2f})"
                    if getattr(region, "dryout_enabled", False) else "YOK"
                ),
                (
                    "PASS" if (not getattr(region, "dryout_enabled", False) or getattr(region, "dryout_converged", False))
                    else "YAKINSAMADI"
                ),
                _status_text(getattr(region, "status", "")),
            ))
            warnings.extend(str(x) for x in getattr(region, "warnings", ()) if str(x).strip())
    tables = [ReportTable(
        "2D nodal güzergâh senaryoları",
        ("Senaryo", "Ad", "I/kablo [A]", "Güzergâh ampacity [A]", "Kritik bölge", "Tcond max [°C]", "Durum"),
        tuple(rows),
    )]
    method_validation = getattr(result, "method_validation", None)
    if method_validation is not None:
        method_rows = []
        method_region_rows = []
        for scenario in getattr(method_validation, "scenarios", ()):
            method_rows.append((
                _status_text(getattr(scenario, "scenario_id", "")),
                _status_text(getattr(scenario, "calculation_basis", "")),
                _status_text(getattr(scenario, "validation_status", "")),
                _status_text(getattr(scenario, "judgement_basis_status", "")),
                _fmt(getattr(scenario, "official_ampacity_a", None), 3),
                _fmt(getattr(scenario, "analytic_ampacity_a", None), 3),
                _fmt(getattr(scenario, "nodal_ampacity_a", None), 3),
                _fmt(getattr(scenario, "ampacity_difference_percent", None), 4),
                _fmt(getattr(scenario, "temperature_difference_c", None), 3),
            ))
            for region in getattr(scenario, "region_comparisons", ()):
                quality = getattr(region, "nodal_quality", None)
                method_region_rows.append((
                    _status_text(getattr(scenario, "scenario_id", "")),
                    _status_text(getattr(region, "region_id", "")),
                    _status_text(getattr(region, "reduction_class", "")),
                    _status_text(getattr(region, "calculation_basis", "")),
                    _status_text(getattr(region, "validation_status", "")),
                    _fmt(getattr(region, "analytic_ampacity_a", None), 3),
                    _fmt(getattr(region, "nodal_ampacity_a", None), 3),
                    _fmt(getattr(region, "ampacity_difference_percent", None), 4),
                    _fmt(getattr(region, "temperature_difference_c", None), 3),
                    _status_text(getattr(quality, "status", "")),
                ))
        tables.append(ReportTable(
            "Analitik–nodal yöntem otoritesi",
            ("Senaryo", "Hesap temeli", "Doğrulama", "Veri hüküm temeli", "Resmî [A]", "Analitik [A]", "Nodal [A]", "ΔI [%]", "ΔT [°C]"),
            tuple(method_rows),
        ))
        if method_region_rows:
            tables.append(ReportTable(
                "Bölgesel yöntem karşılaştırması",
                ("Senaryo", "Bölge", "İndirgeme", "Hesap temeli", "Doğrulama", "Analitik [A]", "Nodal [A]", "ΔI [%]", "ΔT [°C]", "Nodal kalite"),
                tuple(method_region_rows),
            ))
        if getattr(method_validation, "validation_status", "") not in {"PASS", "NOT_APPLICABLE"}:
            warnings.append(
                "Analitik–nodal yöntem otoritesi "
                f"{getattr(method_validation, 'calculation_basis', '')} / "
                f"{getattr(method_validation, 'validation_status', '')}; kesin yöntem hükmü için inceleme gerekebilir."
            )
        trace.extend(str(x) for x in getattr(method_validation, "trace", ()))
    if region_rows:
        tables.append(ReportTable(
            "2D nodal bölgesel sonuçlar",
            ("Senaryo", "Bölge ID", "Bölge", "Başlangıç [m]", "Bitiş [m]", "I/kablo [A]", "Ampacity [A]", "Tcond max [°C]", "Enerji dengesi [%]", "Mesh", "Kuruyan hücre", "Kuruma çözümü", "Durum"),
            tuple(region_rows),
        ))
    return ReportSection(
        MODULE_NODAL,
        MODULE_LABELS[MODULE_NODAL],
        status=_status_text(getattr(getattr(result, "active", None), "status", getattr(result, "status", ""))),
        tables=tuple(tables),
        warnings=tuple(dict.fromkeys(warnings)),
        trace=tuple(trace),
    )


def _section_transient(result: Any | None) -> ReportSection:
    if result is None:
        return _not_run(MODULE_TRANSIENT)
    rows = []
    warnings = []
    trace = list(str(x) for x in getattr(result, "trace", ()))
    for region in getattr(result, "regions", ()):
        rows.append((
            _status_text(getattr(region, "region_id", "")),
            _status_text(getattr(region, "region_name", "")),
            _fmt(getattr(region, "cyclic_rating_per_cable_a", None), 3),
            _fmt(getattr(region, "emergency_rating_per_cable_a", None), 3),
            _fmt(getattr(region, "maximum_conductor_temperature_c", None), 3),
            _fmt(getattr(region, "time_of_maximum_h", None), 3),
            _status_text(getattr(region, "status", "")),
        ))
        warnings.extend(str(x) for x in getattr(region, "warnings", ()) if str(x).strip())
        trace.extend(str(x) for x in getattr(region, "trace", ()))
    summary = (
        ("Profil", _status_text(getattr(result, "profile_name", ""))),
        ("Profil tepe akım çarpanı", _fmt(getattr(result, "peak_profile_multiplier", None), 6)),
        ("Akım yük faktörü LF", _fmt(getattr(result, "current_load_factor", None), 6)),
        ("IEC 60853 kayıp-yük faktörü μ", _fmt(getattr(result, "loss_load_factor_mu", None), 6)),
        ("Güzergâh çevrimsel rating", _fmt(getattr(result, "route_cyclic_rating_per_cable_a", None), 3, " A")),
        ("Kritik çevrimsel bölge", _status_text(getattr(result, "critical_cyclic_region_id", ""))),
        ("Güzergâh acil durum rating", _fmt(getattr(result, "route_emergency_rating_per_cable_a", None), 3, " A")),
        ("Kritik acil durum bölgesi", _status_text(getattr(result, "critical_emergency_region_id", ""))),
        ("Maksimum iletken sıcaklığı", _fmt(getattr(result, "maximum_conductor_temperature_c", None), 3, " °C")),
        ("Durum", _status_text(getattr(result, "status", ""))),
    )
    return ReportSection(
        MODULE_TRANSIENT,
        MODULE_LABELS[MODULE_TRANSIENT],
        status=_status_text(getattr(result, "status", "")),
        tables=(
            ReportTable("IEC 60853 güzergâh özeti", ("Parametre", "Değer"), summary),
            ReportTable("Bölgesel geçici/çevrimsel sonuçlar", ("Bölge ID", "Bölge", "Çevrimsel [A]", "Acil [A]", "Tcond max [°C]", "Tmax zamanı [h]", "Durum"), tuple(rows)),
        ),
        warnings=tuple(dict.fromkeys(warnings)),
        trace=tuple(trace),
    )


def _section_bonding(result: Any | None, project: ProjectData, production: Any | None = None) -> ReportSection:
    if result is None:
        return _not_run(MODULE_BONDING)
    summary = (
        ("Bonding şeması", _status_text(getattr(result, "scheme", ""))),
        ("Çözüm modu", _status_text(getattr(result, "solver_mode", ""))),
        ("Toplam uzunluk", _fmt(getattr(result, "total_length_m", None), 3, " m")),
        ("Major section sayısı", _fmt(getattr(result, "major_section_count", None), 0)),
        ("Maksimum standing voltage", _fmt(getattr(result, "max_standing_voltage_v", None), 3, " V")),
        ("Proje voltage limiti", _fmt(getattr(result, "voltage_limit_v", None), 3, " V")),
        ("Voltage kriteri", "PASS" if getattr(result, "voltage_limit_ok", False) else "FAIL"),
        ("Bonding lead kriteri", "PASS" if getattr(result, "lead_length_ok", False) else "FAIL"),
        ("Toplam metalik kılıf kaybı", _fmt(getattr(result, "total_sheath_loss_w", None), 3, " W")),
        ("Legacy boyuna kılıf kayıp oranı", _fmt(getattr(result, "lambda1", None), 7)),
        ("İdeal iptal göstergesi", _fmt(getattr(result, "ideal_cancellation", None))),
    )
    loop_rows = []
    for loop in getattr(result, "loop_results", ()):
        loop_rows.append((
            _fmt(getattr(loop, "major_index", None), 0),
            _status_text(getattr(loop, "loop_name", "")),
            "→".join(getattr(loop, "sheath_path", ())),
            _fmt(getattr(loop, "residual_emf_magnitude_v", None), 5),
            _fmt(getattr(loop, "current_magnitude_a", None), 5),
            _fmt(getattr(loop, "sheath_loss_w", None), 5),
        ))
    accessory_rows = []
    accessory_plan = resolve_bonding_accessory_plan(project.bonding)
    item_by_box = {item.link_box_id: item for item in accessory_plan.items}
    for box in project.bonding.link_boxes:
        resolved = item_by_box.get(box.link_box_id)
        accessory_rows.append((
            box.link_box_id,
            box.name,
            _fmt(box.position_m, 3),
            _fmt(box.lead_length_m, 3),
            box.lead_type,
            resolved.boundary_role if resolved else "ENGINEERING_REVIEW",
            resolved.link_box_role if resolved else "CUSTOM_LINK_BOX",
            resolved.svl_requirement if resolved else "ENGINEERING_REVIEW",
            box.svl_candidate_id or "—",
        ))
    tables = []
    if production is not None and getattr(production, "scenarios", ()):
        prod_rows = []
        for item in production.scenarios:
            currents = ", ".join(f"{cid}:{amps:.3f} A" for cid, amps in item.circuit_currents_a)
            prod_rows.append((
                _status_text(item.scenario_id),
                currents,
                ",".join(item.deenergized_circuit_ids) or "—",
                _fmt(item.maximum_sheath_current_a, 5),
                _fmt(item.maximum_sheath_to_earth_voltage_v, 5),
                _fmt(item.maximum_sheath_to_sheath_voltage_v, 5),
                _fmt(item.total_sheath_metal_loss_w, 5),
                _fmt(item.lambda1, 8),
                "PASS" if item.converged and item.methods_agree else (item.error_code or "INDETERMINATE"),
            ))
        tables.append(ReportTable(
            "Üretim bonding senaryoları — global N-core/N-kılıf otoritesi",
            ("Senaryo", "Devre akımları", "Devre dışı", "Ish maks [A]", "Vsh-e maks [V]",
             "Vsh-sh maks [V]", "Kılıf kaybı [W]", "Network sheath-loss ratio", "Durum"),
            tuple(prod_rows),
            notes=("Bu tablo bonding ağının üretim otoritesidir; gösterilen oran yalnız boyuna metalik kılıf I²R bileşenidir, toplam IEC λ1 değildir.", "Paralel devrelerin karşılıklı elektromanyetik katkısı aynı global ağda çözülür."),
        ))
    tables.append(ReportTable(
        "Legacy üç-loop bonding özeti — tanısal karşılaştırma, üretim otoritesi değildir",
        ("Parametre", "Değer"), summary,
    ))
    if loop_rows:
        tables.append(ReportTable("Metalik kılıf çevrimleri", ("Major", "Çevrim", "Yol", "Residual EMF [V]", "Akım [A]", "Kayıp [W]"), tuple(loop_rows)))
    if accessory_rows:
        tables.append(ReportTable("Link box yerleşimleri", ("ID", "Ad", "Konum [m]", "Lead [m]", "Lead tipi", "Sınır rolü", "Kutu rolü", "SVL gereksinimi", "SVL adayı"), tuple(accessory_rows)))
    warnings = [str(x) for x in getattr(result, "notes", ()) if str(x).strip()]
    if not getattr(result, "voltage_limit_ok", True):
        warnings.append("Metalik kılıf standing-voltage proje kriteri aşılmıştır.")
    if not getattr(result, "lead_length_ok", True):
        warnings.append("Bir veya daha fazla bonding lead proje uzunluk kriterini aşmıştır.")
    trace = tuple(str(x) for x in result.trace_lines()) if hasattr(result, "trace_lines") else tuple(str(x) for x in getattr(result, "trace", ()))
    return ReportSection(MODULE_BONDING, MODULE_LABELS[MODULE_BONDING], tables=tuple(tables), warnings=tuple(dict.fromkeys(warnings)), trace=trace)


def _section_fault(result: Any | None) -> ReportSection:
    if result is None:
        return _not_run(MODULE_FAULT)
    summary = (
        ("Çözüm yöntemi", _status_text(getattr(result, "selected_method", ""))),
        ("Yönetici senaryo", _status_text(getattr(result, "governing_scenario_name", ""))),
        ("Yönetici TOV", _fmt(getattr(result, "governing_tov_rms_v", None), 3, " V rms")),
        ("Yönetici süre", _fmt(getattr(result, "governing_duration_s", None), 4, " s")),
        ("Maksimum EPR", _fmt(getattr(result, "maximum_epr_v", None), 3, " V")),
        ("Maksimum metalik kılıf akımı", _fmt(getattr(result, "maximum_sheath_current_a", None), 3, " A")),
        ("Maksimum GCC/ECC akımı", _fmt(getattr(result, "maximum_gcc_current_a", None), 3, " A")),
        ("CIM/NV yöntem uyumu", "PASS" if getattr(result, "all_methods_agree", False) else "FAIL"),
    )
    rows = []
    for scenario in getattr(result, "scenario_results", ()):
        rows.append((
            _status_text(getattr(scenario, "scenario_id", "")),
            _status_text(getattr(scenario, "name", "")),
            _status_text(getattr(scenario, "fault_type", "")),
            _fmt(getattr(scenario, "fault_current_a", None), 3),
            _fmt(getattr(scenario, "duration_s", None), 4),
            _fmt(getattr(scenario, "maximum_sheath_current_a", None), 3),
            _fmt(getattr(scenario, "maximum_gcc_current_a", None), 3),
            _fmt(getattr(scenario, "maximum_epr_v", None), 3),
            _fmt(getattr(scenario, "governing_tov_rms_v", None), 3),
            "PASS" if getattr(scenario, "methods_agree", False) else "FAIL",
        ))
    warnings = [str(x) for x in getattr(result, "notes", ()) if str(x).strip()]
    if not getattr(result, "all_methods_agree", True):
        warnings.append("CIM ve Node-Voltage çözümleri kabul toleransında uyuşmamaktadır.")
    trace = tuple(str(x) for x in result.trace_lines()) if hasattr(result, "trace_lines") else ()
    return ReportSection(
        MODULE_FAULT,
        MODULE_LABELS[MODULE_FAULT],
        tables=(
            ReportTable("Arıza/EPR özeti", ("Parametre", "Değer"), summary),
            ReportTable("Arıza senaryoları", ("ID", "Senaryo", "Tip", "If [A]", "Süre [s]", "Ikılıf max [A]", "IGCC max [A]", "EPR max [V]", "TOV [V]", "Yöntem uyumu"), tuple(rows)),
        ),
        warnings=tuple(dict.fromkeys(warnings)),
        trace=trace,
    )


def _section_svl(result: Any | None) -> ReportSection:
    if result is None:
        return _not_run(MODULE_SVL)
    summary = (
        ("Önerilen aday", _status_text(getattr(result, "recommended_display_name", ""))),
        ("Önerilen aday ID", _status_text(getattr(result, "recommended_candidate_id", ""))),
        ("Normal standing voltage", _fmt(getattr(result, "normal_standing_voltage_rms_v", None), 3, " V rms")),
        ("Acil standing voltage", _fmt(getattr(result, "emergency_standing_voltage_rms_v", None), 3, " V rms")),
        ("Gerekli sürekli gerilim", _fmt(getattr(result, "continuous_required_rms_v", None), 3, " V rms")),
        ("En uzun bonding lead", _fmt(getattr(result, "worst_bonding_lead_length_m", None), 3, " m")),
    )
    rows = []
    candidate_notes: list[str] = []
    for check in getattr(result, "checks", ()):
        candidate_id = _status_text(getattr(check, "candidate_id", ""))
        failed = tuple(str(x) for x in getattr(check, "failed_checks", ()) if str(x).strip())
        pending = tuple(str(x) for x in getattr(check, "pending_checks", ()) if str(x).strip())
        rows.append((
            candidate_id,
            _status_text(getattr(check, "display_name", "")),
            _fmt(getattr(check, "mcov_rms_v", None), 3),
            _fmt(getattr(check, "tov_withstand_rms_v", None), 3),
            _fmt(getattr(check, "protective_level_peak_v", None), 3),
            _fmt(getattr(check, "energy_capacity_kj", None), 3),
            _fmt(getattr(check, "nominal_discharge_current_ka", None), 3),
            _status_text(getattr(check, "status", "")),
        ))
        if failed:
            candidate_notes.append(f"{candidate_id} — başarısız kontroller: " + "; ".join(failed))
        if pending:
            candidate_notes.append(f"{candidate_id} — bekleyen kontroller: " + "; ".join(pending))
    warnings = [str(x) for x in getattr(result, "notes", ()) if str(x).strip()]
    if not getattr(result, "recommended_candidate_id", ""):
        warnings.append("SVL için öneri üretilememiştir.")
    trace = tuple(str(x) for x in result.trace_lines()) if hasattr(result, "trace_lines") else tuple(str(x) for x in getattr(result, "trace", ()))
    return ReportSection(
        MODULE_SVL,
        MODULE_LABELS[MODULE_SVL],
        tables=(
            ReportTable("SVL seçim özeti", ("Parametre", "Değer"), summary),
            ReportTable(
                "SVL aday kontrolleri",
                ("ID", "Aday", "MCOV [V]", "TOV [V]", "Koruma seviyesi [Vpk]", "Enerji [kJ]", "In [kA]", "Durum"),
                tuple(rows),
                tuple(candidate_notes),
            ),
        ),
        warnings=tuple(dict.fromkeys(warnings)),
        trace=trace,
    )


def _mandatory_warnings(project: ProjectData, sections: Sequence[ReportSection]) -> tuple[str, ...]:
    warnings: list[str] = []
    audit = audit_project_sources(project)
    gates = evaluate_application_iteration_gates(deepcopy(project))
    if audit.status not in {"PASS", "READY", "OK"}:
        warnings.append(f"Kaynak veri denetimi {audit.status}: {audit.issue_count} bulgu, {audit.missing_data_count} eksik veri.")
    if gates.status != "READY":
        warnings.append(f"Kablo uygulama/iterasyon kapıları {gates.status}; sonuçlar nihai tasarım onayı değildir.")
    if project.design_progress.final_design != "READY":
        warnings.append(f"Proje nihai tasarım durumu: {project.design_progress.final_design}.")
    if project.cable.data_status not in {"VALIDATED", "TESTED", "AS_BUILT"}:
        warnings.append(f"Kablo verisi doğrulanmış/as-built seviyede değildir: {project.cable.data_status}.")
    for section in sections:
        if section.status in {"NOT_RUN", "NOT_AVAILABLE", "BLOCKED", "FAIL"}:
            warnings.append(f"{section.title}: {section.status}.")
        warnings.extend(section.warnings)
    cleaned: list[str] = []
    for item in warnings:
        text = re.sub(r"\s+", " ", str(item)).strip()
        if text and text not in cleaned:
            cleaned.append(text)
    return tuple(cleaned)


def _section_procurement(project: ProjectData) -> ReportSection:
    try:
        package = build_procurement_package(project)
    except Exception as exc:
        return ReportSection(
            MODULE_PROCUREMENT,
            MODULE_LABELS[MODULE_PROCUREMENT],
            status="NOT_AVAILABLE",
            warnings=(f"BOQ/BOM paketi üretilemedi: {exc}",),
        )
    s = package.summary
    summary_rows = (
        ("Net yeraltı güzergâhı", _fmt(s.net_route_length_m, 3, " m")),
        ("Montajlı tek damarlı kablo", _fmt(s.installed_single_core_length_m, 3, " m")),
        ("Sipariş tek damarlı kablo", _fmt(s.order_single_core_length_m, 3, " m")),
        ("Termination", _fmt(s.termination_units, 0, " adet")),
        ("Joint", _fmt(s.joint_units, 0, " adet")),
        ("Cross-bonding link box", _fmt(s.cross_bonding_link_box_units, 0, " adet")),
        ("Grounding link box", _fmt(s.grounding_link_box_units, 0, " adet")),
        ("Toplam link box", _fmt(s.link_box_units, 0, " adet")),
        ("SVL seti", _fmt(s.svl_set_units, 0, " set")),
        ("SVL elemanı/polü", _fmt(s.svl_pole_units, 0, " adet")),
        ("Makara", _fmt(s.drum_count, 0, " adet")),
        ("Makara planı", s.drum_plan_status),
        ("Sipariş/fire payı", _fmt(s.order_allowance_total_m, 3, " m")),
        ("Toplam makara aşımı", _fmt(s.overload_total_m, 3, " m")),
        ("Tedarik paketi durumu", s.status),
    )
    item_rows = tuple(
        (item.item_id, item.description, _fmt(item.final_quantity, 3), item.unit, item.status)
        for item in package.lines
    )
    return ReportSection(
        MODULE_PROCUREMENT,
        MODULE_LABELS[MODULE_PROCUREMENT],
        status=s.status,
        paragraphs=(
            "Bu bölüm proje modelinden türetilen tedarik özetidir. Ayrıntılı BOQ, BOM, RFQ ve makara planı ayrı tedarik çıktılarında üretilir.",
        ),
        tables=(
            ReportTable("Tedarik ve metraj özeti", ("Gösterge", "Değer"), summary_rows),
            ReportTable("Ana BOQ/BOM kalemleri", ("Kalem", "Tanım", "Miktar", "Birim", "Durum"), item_rows),
        ),
        warnings=package.warnings,
        trace=(package.reference,) + package.assumptions,
    )


def build_project_report(
    project: ProjectData,
    configuration: ReportConfiguration | None = None,
    results: CalculationResultsBundle | None = None,
) -> ProjectReport:
    config = configuration or default_report_configuration()
    bundle = results or CalculationResultsBundle()
    modules = config.resolved_modules()
    builders = {
        MODULE_PROJECT: lambda: _section_project(project),
        MODULE_DESIGN_BASIS: lambda: _section_design_basis(project),
        MODULE_SOURCE_AUDIT: lambda: _section_source_audit(project),
        MODULE_CABLE: lambda: _section_cable(project),
        MODULE_ROUTE: lambda: _section_route(project),
        MODULE_VOLTAGE_DROP: lambda: _section_voltage_drop(project),
        MODULE_IEC60287: lambda: _section_iec(bundle.iec_results, bundle.production_electrothermal_result),
        MODULE_NODAL: lambda: _section_nodal(bundle.nodal_thermal_result),
        MODULE_TRANSIENT: lambda: _section_transient(bundle.transient_thermal_result),
        MODULE_BONDING: lambda: _section_bonding(bundle.bonding_result, project, bundle.production_bonding_result),
        MODULE_FAULT: lambda: _section_fault(bundle.fault_result),
        MODULE_SVL: lambda: _section_svl(bundle.svl_result),
        MODULE_PROCUREMENT: lambda: _section_procurement(project),
    }
    sections: list[ReportSection] = []
    calculation_trace: list[str] = []
    trace_counts: list[tuple[str, int]] = []
    for module in modules:
        if module in {MODULE_WARNINGS, MODULE_TRACE}:
            continue
        full_section = builders[module]()
        trace_counts.append((full_section.title, len(full_section.trace)))
        calculation_trace.extend(full_section.trace)
        section = full_section if config.include_detailed_trace else replace(full_section, trace=())
        if config.include_empty_selected_modules or section.status not in {"NOT_RUN", "NOT_AVAILABLE"}:
            sections.append(section)

    mandatory = _mandatory_warnings(project, sections)
    if MODULE_TRACE in modules:
        trace_lines = calculation_trace if config.include_detailed_trace else [
            f"{title}: {count} hesap izi satırı mevcut."
            for title, count in trace_counts if count
        ]
        sections.append(ReportSection(
            MODULE_TRACE,
            MODULE_LABELS[MODULE_TRACE],
            paragraphs=(
                "Ayrıntılı iz seçilmediyse bu bölüm yalnız iz kapsamını gösterir. Hesap sonuçları, kaynak girdiler ve yazılım sürümü birlikte arşivlenmelidir.",
            ),
            trace=tuple(trace_lines),
        ))
    if MODULE_WARNINGS in modules:
        sections.append(ReportSection(
            MODULE_WARNINGS,
            MODULE_LABELS[MODULE_WARNINGS],
            status="ATTENTION" if mandatory else "CLEAR",
            paragraphs=(
                "Bu bölüm kullanıcı tarafından rapordan çıkarılamaz. Kritik eksik veriler ve koşullu hesap durumları burada görünür tutulur.",
            ),
            warnings=mandatory or ("Kritik uyarı kaydedilmemiştir.",),
        ))

    title = config.metadata.title.strip() or REPORT_TYPE_LABELS.get(
        config.metadata.report_type, "DiTuS Teknik Rapor"
    )
    metadata = ReportMetadata(**{**asdict(config.metadata), "title": title})
    blocked = any(section.status in {"BLOCKED", "FAIL", "NOT_AVAILABLE"} for section in sections)
    not_run = any(section.status == "NOT_RUN" for section in sections)
    report_status = "CONDITIONAL" if mandatory or blocked or not_run else "READY"
    trace = (
        REFERENCE,
        f"Rapor türü: {metadata.report_type}",
        f"Seçilen modüller: {', '.join(modules)}",
        f"Proje imzası: {_stable_project_signature(project)}",
        "Uyarılar ve sınırlamalar bölümü zorunlu olarak eklendi.",
    )
    return ProjectReport(
        reference=REFERENCE,
        generated_at=datetime.now().isoformat(timespec="seconds"),
        project_name=project.project_name,
        project_code=project.project_code,
        project_signature_sha256=_stable_project_signature(project),
        metadata=metadata,
        selected_modules=modules,
        sections=tuple(sections),
        mandatory_warnings=mandatory,
        report_status=report_status,
        trace=trace,
    )


def render_project_report_markdown(report: ProjectReport) -> str:
    m = report.metadata
    lines = [
        f"# {m.title}",
        "",
        f"**Proje:** {report.project_name}  ",
        f"**Proje kodu:** {report.project_code}  ",
        f"**Doküman no:** {m.document_no or '—'}  ",
        f"**Revizyon:** {m.revision or '—'}  ",
        f"**Yayın tarihi:** {m.issue_date or '—'}  ",
        f"**Hazırlayan / kontrol:** {m.prepared_by or '—'} / {m.checked_by or '—'}  ",
        f"**Durum:** {m.approval_status} / {report.report_status}  ",
        f"**Proje veri imzası:** `{report.project_signature_sha256}`",
        "",
        "> Bu rapor nihai uygunluk kararı değildir; eksik veya koşullu girdileri uygunluğa dönüştürmez. Uyarılar ve sınırlamalar bölümü bağlayıcıdır.",
        "",
    ]
    for index, section in enumerate(report.sections, start=1):
        lines.extend([f"## {index}. {section.title}", "", f"**Bölüm durumu:** `{section.status}`", ""])
        for paragraph in section.paragraphs:
            lines.extend([paragraph, ""])
        for table in section.tables:
            lines.extend([f"### {table.title}", ""])
            if table.headers:
                lines.append("| " + " | ".join(table.headers) + " |")
                lines.append("| " + " | ".join("---" for _ in table.headers) + " |")
                for row in table.rows:
                    lines.append("| " + " | ".join(str(cell).replace("|", "\\|") for cell in row) + " |")
                lines.append("")
            for note in table.notes:
                lines.extend([f"- {note}", ""])
        if section.warnings:
            lines.extend(["### Uyarılar", ""])
            lines.extend(f"- {warning}" for warning in section.warnings)
            lines.append("")
        if section.trace:
            lines.extend(["### Hesap izi", "", "```text"])
            lines.extend(section.trace)
            lines.extend(["```", ""])
    lines.extend([
        "---",
        f"Rapor üretim zamanı: {report.generated_at}",
        f"Rapor motoru: {report.reference}",
    ])
    return "\n".join(lines).rstrip() + "\n"


def render_project_report_html(report: ProjectReport) -> str:
    m = report.metadata
    section_html: list[str] = []
    for index, section in enumerate(report.sections, start=1):
        tables = []
        for table in section.tables:
            head = "".join(f"<th>{escape(str(cell))}</th>" for cell in table.headers)
            body = "".join(
                "<tr>" + "".join(f"<td>{escape(str(cell))}</td>" for cell in row) + "</tr>"
                for row in table.rows
            )
            notes = "".join(f"<li>{escape(note)}</li>" for note in table.notes)
            tables.append(
                f"<h3>{escape(table.title)}</h3><div class='table-wrap'><table><thead><tr>{head}</tr></thead>"
                f"<tbody>{body}</tbody></table></div>" + (f"<ul>{notes}</ul>" if notes else "")
            )
        paragraphs = "".join(f"<p>{escape(p)}</p>" for p in section.paragraphs)
        warnings = "".join(f"<li>{escape(w)}</li>" for w in section.warnings)
        trace = "\n".join(escape(line) for line in section.trace)
        section_html.append(
            f"<section><div class='section-heading'><span>{index}</span><h2>{escape(section.title)}</h2>"
            f"<b class='status'>{escape(section.status)}</b></div>{paragraphs}{''.join(tables)}"
            + (f"<div class='warning'><h3>Uyarılar</h3><ul>{warnings}</ul></div>" if warnings else "")
            + (f"<details><summary>Hesap izi</summary><pre>{trace}</pre></details>" if trace else "")
            + "</section>"
        )
    return f"""<!doctype html>
<html lang='tr'><head><meta charset='utf-8'><title>{escape(m.title)}</title>
<style>
@page {{ size:A4; margin:16mm 14mm 16mm 14mm; }}
:root {{ --ink:#17324a; --muted:#5c6b77; --line:#cbd5de; --soft:#eef3f7; --warn:#fff2cf; --warnline:#c98900; }}
* {{ box-sizing:border-box; }} body {{ margin:0; font-family:Arial,'Noto Sans',sans-serif; color:#1f2933; background:#f4f6f8; line-height:1.42; }}
.report {{ width:min(1100px,100%); margin:24px auto; background:white; box-shadow:0 5px 22px #0002; }}
.cover {{ min-height:920px; padding:64px 68px; display:flex; flex-direction:column; border-top:12px solid var(--ink); }}
.brand {{ font-size:18px; font-weight:700; color:var(--ink); letter-spacing:.03em; }}
.cover h1 {{ font-size:34px; line-height:1.12; margin:110px 0 16px; color:var(--ink); }}
.subtitle {{ color:var(--muted); font-size:18px; }}
.meta {{ margin-top:auto; display:grid; grid-template-columns:180px 1fr; border-top:1px solid var(--line); }}
.meta div {{ padding:8px 10px; border-bottom:1px solid var(--line); }} .meta div:nth-child(odd) {{ background:var(--soft); font-weight:700; }}
.content {{ padding:42px 54px 70px; }} section {{ margin:0 0 36px; page-break-inside:avoid; }}
.section-heading {{ display:flex; align-items:center; gap:12px; background:var(--ink); color:#fff; padding:8px 11px; margin-bottom:14px; }}
.section-heading span {{ width:34px; height:34px; border-radius:50%; background:#fff; color:var(--ink); display:grid; place-items:center; font-weight:700; }}
h2 {{ color:#fff; font-size:22px; flex:1; margin:0; }} h3 {{ color:#294e69; font-size:15px; margin:18px 0 7px; }}
.status {{ color:#fff; font-size:11px; }} .table-wrap {{ overflow-x:auto; }}
table {{ width:100%; border-collapse:collapse; font-size:12px; }} th {{ background:var(--ink); color:white; text-align:left; }} th,td {{ border:1px solid var(--line); padding:6px 7px; vertical-align:top; }} tbody tr:nth-child(even) {{ background:#f8fafb; }}
.warning {{ border-left:5px solid var(--warnline); background:var(--warn); padding:9px 14px; margin-top:14px; }} .warning h3 {{ margin-top:0; color:#7a4d00; }}
pre {{ white-space:pre-wrap; background:#101a23; color:#e6edf3; padding:12px; font-size:11px; }}
.footer-note {{ margin-top:42px; color:var(--muted); font-size:11px; border-top:1px solid var(--line); padding-top:10px; }}
@media print {{ body {{ background:white; }} .report {{ margin:0; width:auto; box-shadow:none; }} .cover {{ page-break-after:always; min-height:260mm; }} details {{ display:block; }} details summary {{ display:none; }} section {{ page-break-inside:auto; }} tr {{ page-break-inside:avoid; }} }}
</style></head><body><main class='report'>
<div class='cover'><div class='brand'>DiTuS Kablo Analizör™</div><h1>{escape(m.title)}</h1>
<div class='subtitle'>{escape(report.project_name)} · {escape(report.project_code)}</div>
<div class='meta'>
<div>Doküman No</div><div>{escape(m.document_no or '—')}</div><div>Revizyon</div><div>{escape(m.revision or '—')}</div>
<div>Yayın Tarihi</div><div>{escape(m.issue_date or '—')}</div><div>İşveren</div><div>{escape(m.client or '—')}</div>
<div>Yüklenici</div><div>{escape(m.contractor or '—')}</div><div>Hazırlayan</div><div>{escape(m.prepared_by or '—')}</div>
<div>Kontrol Eden</div><div>{escape(m.checked_by or '—')}</div><div>Durum</div><div>{escape(m.approval_status)} / {escape(report.report_status)}</div>
<div>Proje İmzası</div><div><code>{escape(report.project_signature_sha256)}</code></div>
</div></div><div class='content'>
<div class='warning'><strong>Rapor sınırlaması:</strong> Eksik veya koşullu girdiler nihai uygunluk olarak yorumlanamaz. Uyarılar ve sınırlamalar bölümü raporun ayrılmaz parçasıdır.</div>
{''.join(section_html)}
<div class='footer-note'>Üretim: {escape(report.generated_at)} · {escape(report.reference)}</div></div></main></body></html>"""


def _write_docx(report: ProjectReport, path: Path) -> None:
    try:
        from docx import Document
        from docx.enum.section import WD_SECTION
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt, RGBColor
    except ImportError as exc:
        raise ReportGenerationError("DOCX çıktısı için python-docx kurulmalıdır.") from exc

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.5)
    styles = doc.styles
    styles["Normal"].font.name = "Liberation Sans"
    styles["Normal"].font.size = Pt(9)
    for style_name, size in (("Title", 27), ("Heading 1", 17), ("Heading 2", 13), ("Heading 3", 10)):
        styles[style_name].font.name = "Liberation Sans"
        styles[style_name].font.size = Pt(size)

    dark_fill = "17324A"
    light_text = RGBColor(255, 255, 255)

    def _shade_paragraph(paragraph, fill: str, font_color: RGBColor) -> None:
        p_pr = paragraph._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        p_pr.append(shd)
        for run in paragraph.runs:
            run.font.color.rgb = font_color

    def _shade_cell(cell, fill: str, font_color: RGBColor, bold: bool = True) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = bold
                run.font.color.rgb = font_color

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DiTuS Kablo Analizör™")
    run.bold = True
    run.font.size = Pt(16)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(report.metadata.title)
    run.bold = True
    run.font.size = Pt(28)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"{report.project_name}\n{report.project_code}").font.size = Pt(14)
    doc.add_paragraph("")
    cover = doc.add_table(rows=0, cols=2)
    cover.style = "Table Grid"
    cover.autofit = True
    cover_data = (
        ("Doküman No", report.metadata.document_no or "—"),
        ("Revizyon", report.metadata.revision or "—"),
        ("Yayın Tarihi", report.metadata.issue_date or "—"),
        ("İşveren", report.metadata.client or "—"),
        ("Yüklenici", report.metadata.contractor or "—"),
        ("Hazırlayan", report.metadata.prepared_by or "—"),
        ("Kontrol Eden", report.metadata.checked_by or "—"),
        ("Durum", f"{report.metadata.approval_status} / {report.report_status}"),
        ("Proje İmzası", report.project_signature_sha256),
    )
    for label, value in cover_data:
        cells = cover.add_row().cells
        cells[0].text = label
        cells[1].text = value
        for run in cells[0].paragraphs[0].runs:
            run.bold = True
    doc.add_paragraph("")
    limitation = doc.add_paragraph()
    limitation.style = doc.styles["Normal"]
    run = limitation.add_run("Rapor sınırlaması: ")
    run.bold = True
    limitation.add_run("Bu rapor nihai uygunluk kararı değildir. Eksik veya koşullu girdiler uygunluğa dönüştürülmez; uyarılar ve sınırlamalar bölümü bağlayıcıdır.")
    doc.add_page_break()

    for index, report_section in enumerate(report.sections, start=1):
        heading = doc.add_heading(f"{index}. {report_section.title}", level=1)
        _shade_paragraph(heading, dark_fill, light_text)
        status = doc.add_paragraph()
        status.add_run("Bölüm durumu: ").bold = True
        status.add_run(report_section.status)
        for paragraph in report_section.paragraphs:
            doc.add_paragraph(paragraph)
        for table_model in report_section.tables:
            doc.add_heading(table_model.title, level=2)
            table = doc.add_table(rows=1, cols=len(table_model.headers))
            table.style = "Table Grid"
            table.autofit = True
            hdr = table.rows[0].cells
            for i, value in enumerate(table_model.headers):
                hdr[i].text = str(value)
                _shade_cell(hdr[i], dark_fill, light_text, True)
            for row in table_model.rows:
                cells = table.add_row().cells
                for i, value in enumerate(row):
                    cells[i].text = str(value)
            for note in table_model.notes:
                doc.add_paragraph(note, style="List Bullet")
        if report_section.warnings:
            doc.add_heading("Uyarılar", level=2)
            for warning in report_section.warnings:
                doc.add_paragraph(warning, style="List Bullet")
        if report_section.trace:
            doc.add_heading("Hesap izi", level=2)
            for line in report_section.trace:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(0)
                r = p.add_run(line)
                r.font.name = "Liberation Mono"
                r.font.size = Pt(7.5)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(f"DiTuS Kablo Analizör™ · {report.project_code} · {report.metadata.document_no or '—'} · Rev {report.metadata.revision}")
    doc.core_properties.title = report.metadata.title
    doc.core_properties.subject = report.project_name
    doc.core_properties.comments = "Generated by DiTuS Kablo Analizör reporting engine"
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def _write_pdf(report: ProjectReport, path: Path) -> None:
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.platypus import (
            BaseDocTemplate,
            Frame,
            PageTemplate,
            Paragraph,
            PageBreak,
            Spacer,
            Table,
            TableStyle,
            KeepTogether,
        )
    except ImportError as exc:
        raise ReportGenerationError("PDF çıktısı için reportlab kurulmalıdır.") from exc

    font_regular = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
    font_bold = "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"
    if Path(font_regular).exists() and Path(font_bold).exists():
        pdfmetrics.registerFont(TTFont("DiTuS", font_regular))
        pdfmetrics.registerFont(TTFont("DiTuS-Bold", font_bold))
        normal_font, bold_font = "DiTuS", "DiTuS-Bold"
    else:
        normal_font, bold_font = "Helvetica", "Helvetica-Bold"

    page_w, page_h = A4
    doc = BaseDocTemplate(str(path), pagesize=A4, leftMargin=14*mm, rightMargin=14*mm, topMargin=16*mm, bottomMargin=17*mm)
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="normal")

    def on_page(canvas, _doc):
        canvas.saveState()
        canvas.setFont(normal_font, 7)
        canvas.setFillColor(colors.HexColor("#5c6b77"))
        canvas.drawString(doc.leftMargin, 8*mm, f"DiTuS Kablo Analizör™ · {report.project_code} · Rev {report.metadata.revision}")
        canvas.drawRightString(page_w-doc.rightMargin, 8*mm, f"Sayfa {_doc.page}")
        canvas.restoreState()

    doc.addPageTemplates(PageTemplate(id="main", frames=frame, onPage=on_page))
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("TitleD", parent=styles["Title"], fontName=bold_font, fontSize=25, leading=29, textColor=colors.HexColor("#17324a"), alignment=TA_CENTER, spaceAfter=12)
    subtitle_style = ParagraphStyle("SubtitleD", parent=styles["Normal"], fontName=normal_font, fontSize=12, leading=16, textColor=colors.HexColor("#5c6b77"), alignment=TA_CENTER)
    h1 = ParagraphStyle("H1D", parent=styles["Heading1"], fontName=bold_font, fontSize=15, leading=18, textColor=colors.HexColor("#17324a"), spaceBefore=8, spaceAfter=8)
    h2 = ParagraphStyle("H2D", parent=styles["Heading2"], fontName=bold_font, fontSize=10, leading=12, textColor=colors.HexColor("#294e69"), spaceBefore=8, spaceAfter=5)
    body = ParagraphStyle("BodyD", parent=styles["BodyText"], fontName=normal_font, fontSize=8.5, leading=11, spaceAfter=5)
    small = ParagraphStyle("SmallD", parent=body, fontSize=6.7, leading=8.2)
    table_header = ParagraphStyle("TableHeaderD", parent=small, fontName=bold_font, textColor=colors.white)
    warning_style = ParagraphStyle("WarnD", parent=body, backColor=colors.HexColor("#fff2cf"), borderColor=colors.HexColor("#c98900"), borderWidth=0.7, borderPadding=6, leftIndent=2, rightIndent=2)

    story = [Spacer(1, 25*mm), Paragraph("DiTuS Kablo Analizör™", subtitle_style), Spacer(1, 28*mm), Paragraph(escape(report.metadata.title), title_style), Paragraph(f"{escape(report.project_name)}<br/>{escape(report.project_code)}", subtitle_style), Spacer(1, 20*mm)]
    cover_data = [
        [Paragraph(f"<b>{escape(label)}</b>", body), Paragraph(escape(value), body)]
        for label, value in (
            ("Doküman No", report.metadata.document_no or "—"),
            ("Revizyon", report.metadata.revision or "—"),
            ("Yayın Tarihi", report.metadata.issue_date or "—"),
            ("İşveren", report.metadata.client or "—"),
            ("Yüklenici", report.metadata.contractor or "—"),
            ("Hazırlayan", report.metadata.prepared_by or "—"),
            ("Kontrol Eden", report.metadata.checked_by or "—"),
            ("Durum", f"{report.metadata.approval_status} / {report.report_status}"),
            ("Proje İmzası", report.project_signature_sha256),
        )
    ]
    cover_table = Table(cover_data, colWidths=[45*mm, 125*mm], repeatRows=0)
    cover_table.setStyle(TableStyle([
        ("GRID", (0,0), (-1,-1), 0.4, colors.HexColor("#cbd5de")),
        ("BACKGROUND", (0,0), (0,-1), colors.HexColor("#eef3f7")),
        ("VALIGN", (0,0), (-1,-1), "TOP"),
        ("LEFTPADDING", (0,0), (-1,-1), 5), ("RIGHTPADDING", (0,0), (-1,-1), 5),
        ("TOPPADDING", (0,0), (-1,-1), 4), ("BOTTOMPADDING", (0,0), (-1,-1), 4),
    ]))
    story.extend([cover_table, Spacer(1, 12*mm), Paragraph("<b>Rapor sınırlaması:</b> Bu rapor nihai uygunluk kararı değildir. Eksik veya koşullu girdiler uygunluğa dönüştürülmez; uyarılar ve sınırlamalar bölümü bağlayıcıdır.", warning_style), PageBreak()])

    for index, section_model in enumerate(report.sections, start=1):
        heading_table = Table([[Paragraph(f"{index}. {escape(section_model.title)}", ParagraphStyle(
            f"H1White{index}", parent=h1, textColor=colors.white, spaceBefore=0, spaceAfter=0
        ))]], colWidths=[doc.width])
        heading_table.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,-1), colors.HexColor("#17324a")),
            ("LEFTPADDING", (0,0), (-1,-1), 7), ("RIGHTPADDING", (0,0), (-1,-1), 7),
            ("TOPPADDING", (0,0), (-1,-1), 5), ("BOTTOMPADDING", (0,0), (-1,-1), 5),
        ]))
        story.append(heading_table)
        story.append(Paragraph(f"<b>Bölüm durumu:</b> {escape(section_model.status)}", body))
        for paragraph in section_model.paragraphs:
            story.append(Paragraph(escape(paragraph), body))
        for table_model in section_model.tables:
            story.append(Paragraph(escape(table_model.title), h2))
            data = [[Paragraph(escape(str(v)), table_header) for v in table_model.headers]]
            for row in table_model.rows:
                data.append([Paragraph(escape(str(v)), small) for v in row])
            column_count = max(1, len(table_model.headers))
            widths = [doc.width / column_count] * column_count
            table = Table(data, colWidths=widths, repeatRows=1, splitByRow=True)
            table.setStyle(TableStyle([
                ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#17324a")),
                ("TEXTCOLOR", (0,0), (-1,0), colors.white),
                ("GRID", (0,0), (-1,-1), 0.3, colors.HexColor("#cbd5de")),
                ("VALIGN", (0,0), (-1,-1), "TOP"),
                ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, colors.HexColor("#f8fafb")]),
                ("LEFTPADDING", (0,0), (-1,-1), 3), ("RIGHTPADDING", (0,0), (-1,-1), 3),
                ("TOPPADDING", (0,0), (-1,-1), 2.5), ("BOTTOMPADDING", (0,0), (-1,-1), 2.5),
            ]))
            story.append(table)
            story.append(Spacer(1, 2*mm))
        if section_model.warnings:
            warning_text = "<br/>".join("• " + escape(item) for item in section_model.warnings)
            story.append(KeepTogether([Paragraph("Uyarılar", h2), Paragraph(warning_text, warning_style)]))
        if section_model.trace:
            story.append(Paragraph("Hesap izi", h2))
            for line in section_model.trace:
                story.append(Paragraph(escape(line), small))
        story.append(Spacer(1, 5*mm))
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.build(story)


def write_project_report(
    report: ProjectReport,
    output_directory: str | Path,
    base_name: str | None = None,
    formats: Iterable[str] | None = None,
) -> dict[str, Path]:
    directory = Path(output_directory)
    directory.mkdir(parents=True, exist_ok=True)
    safe_default = re.sub(r"[^A-Za-z0-9_.-]+", "_", f"{report.project_code}_{report.metadata.report_type}_v{__version__}").strip("_")
    stem = re.sub(r"[^A-Za-z0-9_.-]+", "_", base_name or safe_default).strip("_") or "ditus_report"
    requested = tuple(dict.fromkeys(str(item).lower().lstrip(".") for item in (formats or ("html", "markdown", "json", "docx", "pdf"))))
    paths: dict[str, Path] = {}
    for fmt in requested:
        if fmt == "json":
            path = directory / f"{stem}.json"
            path.write_text(json.dumps(report.to_dict(), ensure_ascii=False, indent=2), encoding="utf-8")
        elif fmt in {"markdown", "md"}:
            path = directory / f"{stem}.md"
            path.write_text(render_project_report_markdown(report), encoding="utf-8")
            fmt = "markdown"
        elif fmt == "html":
            path = directory / f"{stem}.html"
            path.write_text(render_project_report_html(report), encoding="utf-8")
        elif fmt == "docx":
            path = directory / f"{stem}.docx"
            _write_docx(report, path)
        elif fmt == "pdf":
            path = directory / f"{stem}.pdf"
            _write_pdf(report, path)
        else:
            raise ReportGenerationError(f"Desteklenmeyen rapor formatı: {fmt}")
        paths[fmt] = path
    return paths
