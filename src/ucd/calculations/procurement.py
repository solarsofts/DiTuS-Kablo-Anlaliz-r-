from __future__ import annotations

from dataclasses import asdict, dataclass, field
from datetime import datetime
from hashlib import sha256
from html import escape
import csv
import io
import json
import math
from pathlib import Path
import re
from typing import Any, Iterable

from ucd import __version__
from ucd.calculations.installation_coupling import project_with_synchronized_installation_geometry
from ucd.calculations.bonding_accessories import (
    ACCESSORY_VALID,
    BondingAccessoryPlan,
    resolve_bonding_accessory_plan,
)
from ucd.models.project import (
    ProjectData,
    ProcurementData,
    ProcurementQuantityOverride,
    THERMAL_INSTALL_DUCT_BANK,
    THERMAL_INSTALL_HDD,
)

REFERENCE = (
    f"DiTuS v{__version__} BOQ/BOM/RFQ and drum-planning engine. Quantities are derived from "
    "project route, project-assigned cable record, bonding graph and explicit procurement assumptions. "
    "Automatic quantities and user overrides remain separately traceable."
)

STATUS_CONFIRMED = "CONFIRMED_PROJECT_DATA"
STATUS_CONDITIONAL = "CONDITIONAL_PROJECT_DATA"
STATUS_ASSUMPTION = "ENGINEERING_ASSUMPTION"
STATUS_NOT_APPLICABLE = "NOT_APPLICABLE"

CATEGORY_CABLE = "CABLE"
CATEGORY_ACCESSORY = "CABLE_ACCESSORY"
CATEGORY_BONDING = "BONDING_AND_SVL"
CATEGORY_GROUNDING = "GROUNDING"
CATEGORY_CIVIL = "CIVIL_WORKS"
CATEGORY_MARKING = "MARKING_AND_SUPPORT"

VIEW_BOQ = "BOQ"
VIEW_BOM = "BOM"
VIEW_RFQ = "RFQ"

PHASES = ("A", "B", "C")


@dataclass(frozen=True)
class QuantityBasis:
    source_object: str
    route_reference: str = ""
    formula: str = ""
    data_state: str = STATUS_CONFIRMED
    source_reference: str = ""
    notes: str = ""


@dataclass(frozen=True)
class ProcurementLine:
    item_id: str
    category: str
    description: str
    technical_specification: str
    auto_quantity: float
    final_quantity: float
    unit: str
    status: str
    basis: QuantityBasis
    override_rationale: str = ""
    manufacturer: str = ""
    model: str = ""
    required_documents: tuple[str, ...] = ()
    supplier_response_fields: tuple[str, ...] = (
        "Teklif edilen marka/model",
        "Teknik uygunluk",
        "Birim fiyat",
        "Teslim süresi",
    )


@dataclass(frozen=True)
class DrumCut:
    cut_id: str
    route_reference: str
    circuit_no: int
    phase: str
    segment_no: int
    segment_length_m: float
    end_allowance_m: float
    installation_allowance_m: float
    required_cut_length_m: float
    cut_type: str = "ROUTE_SEGMENT"
    notes: str = ""


@dataclass(frozen=True)
class DrumAssignment:
    drum_id: str
    cable_item_id: str
    maximum_length_m: float
    loaded_length_m: float
    remaining_length_m: float
    cuts: tuple[DrumCut, ...]
    route_cut_length_m: float = 0.0
    order_allowance_m: float = 0.0
    spare_stock_length_m: float = 0.0
    rounding_reconciliation_m: float = 0.0
    capacity_balance_m: float = 0.0
    remaining_capacity_m: float = 0.0
    overload_m: float = 0.0
    estimated_net_mass_kg: float = 0.0
    assignment_status: str = "VALID"
    status: str = STATUS_CONFIRMED
    notes: str = ""


@dataclass(frozen=True)
class UnassignedDrumCut:
    cut_id: str
    required_cut_length_m: float
    maximum_length_m: float
    deficit_m: float
    reason: str
    route_reference: str = ""


@dataclass(frozen=True)
class DrumPlanSummary:
    route_cut_total_m: float
    order_allowance_total_m: float
    spare_stock_total_m: float
    rounding_reconciliation_total_m: float
    allocated_total_m: float
    unallocated_total_m: float
    overload_total_m: float
    maximum_single_drum_overload_m: float
    accounting_residual_m: float
    valid_drum_count: int
    invalid_drum_count: int
    unassigned_cut_count: int
    drum_plan_status: str


@dataclass(frozen=True)
class ProcurementSummary:
    net_route_length_m: float
    installed_single_core_length_m: float
    order_single_core_length_m: float
    circuit_count: int
    parallel_cables_per_phase: int
    termination_units: int
    joint_units: int
    link_box_units: int
    svl_units: int
    drum_count: int
    status: str
    cross_bonding_link_box_units: int = 0
    grounding_link_box_units: int = 0
    termination_link_box_units: int = 0
    custom_link_box_units: int = 0
    svl_set_units: int = 0
    svl_pole_units: int = 0
    accessory_plan_status: str = "INCOMPLETE"
    drum_plan_status: str = "INCOMPLETE"
    route_cut_total_m: float = 0.0
    order_allowance_total_m: float = 0.0
    spare_stock_total_m: float = 0.0
    allocated_total_m: float = 0.0
    unallocated_total_m: float = 0.0
    overload_total_m: float = 0.0
    accounting_residual_m: float = 0.0
    invalid_drum_count: int = 0
    unassigned_cut_count: int = 0


@dataclass(frozen=True)
class ProcurementPackage:
    reference: str
    generated_at: str
    project_name: str
    project_code: str
    project_signature_sha256: str
    settings: ProcurementData
    summary: ProcurementSummary
    lines: tuple[ProcurementLine, ...]
    drums: tuple[DrumAssignment, ...]
    unassigned_route_cuts: tuple[UnassignedDrumCut, ...]
    drum_plan: DrumPlanSummary
    accessory_plan: BondingAccessoryPlan
    warnings: tuple[str, ...]
    assumptions: tuple[str, ...]

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)

    def lines_for_view(self, view: str) -> tuple[ProcurementLine, ...]:
        view = str(view).upper()
        if view == VIEW_BOQ:
            return self.lines
        if view == VIEW_BOM:
            return tuple(
                line for line in self.lines
                if line.category not in {CATEGORY_CIVIL} or line.final_quantity > 0
            )
        if view == VIEW_RFQ:
            return tuple(line for line in self.lines if line.final_quantity > 0)
        raise ValueError(f"Bilinmeyen tedarik görünümü: {view}")


class ProcurementInputError(ValueError):
    pass


def _stable_project_signature(project: ProjectData) -> str:
    payload = asdict(project)
    payload.pop("modified_at", None)
    raw = json.dumps(payload, ensure_ascii=False, sort_keys=True, separators=(",", ":"), default=str)
    return sha256(raw.encode("utf-8")).hexdigest()


def _round_up(value: float, increment: float) -> float:
    value = max(0.0, float(value))
    increment = max(1.0e-9, float(increment))
    return math.ceil(value / increment - 1.0e-12) * increment


def _fmt(value: float, decimals: int = 3) -> str:
    if not math.isfinite(float(value)):
        return "—"
    return f"{float(value):.{decimals}f}"


def _active_routes(project: ProjectData) -> list[Any]:
    active_names = {
        item.route_section_name
        for item in project.cable_application.assignments
        if item.active and item.route_section_name
    }
    if not active_names:
        return list(project.route_sections)
    selected = [section for section in project.route_sections if section.name in active_names]
    return selected or list(project.route_sections)


def _catalog_record(project: ProjectData) -> Any | None:
    record_id = project.cable.catalog_record_id or project.cable_application.selected_catalog_record_id
    for record in project.cable_library.records:
        if record.record_id == record_id:
            return record
    return None


def _catalog_delivery_and_mass(project: ProjectData, settings: ProcurementData) -> tuple[float, float, str]:
    record = _catalog_record(project)
    catalog_delivery = 0.0
    mass_kg_km = 0.0
    source = "Proje tedarik ayarı"
    if record is not None:
        try:
            catalog_delivery = float(record.catalog_dimensions.get("delivery_length_m", 0.0) or 0.0)
            mass_kg_km = float(record.catalog_dimensions.get("net_weight_kg_km", 0.0) or 0.0)
            if catalog_delivery > 0:
                source = f"Katalog kaydı {record.record_id}"
        except (TypeError, ValueError):
            pass
    maximum = catalog_delivery if catalog_delivery > 0 else float(settings.maximum_drum_length_m)
    if settings.maximum_drum_length_m > 0:
        maximum = min(maximum, float(settings.maximum_drum_length_m)) if maximum > 0 else float(settings.maximum_drum_length_m)
    return max(1.0, maximum), max(0.0, mass_kg_km), source


def _route_segments(project: ProjectData, route_length_m: float) -> list[float]:
    lengths = [float(item.length_m) for item in project.bonding.minor_sections if float(item.length_m) > 0]
    if lengths and abs(sum(lengths) - route_length_m) <= max(1.0, route_length_m * 0.002):
        return lengths
    active = _active_routes(project)
    lengths = [float(item.length_m) for item in active if float(item.length_m) > 0]
    return lengths or [route_length_m]


def _override_map(settings: ProcurementData) -> dict[str, ProcurementQuantityOverride]:
    return {item.item_id: item for item in settings.quantity_overrides if item.item_id}


def _line(
    settings: ProcurementData,
    item_id: str,
    category: str,
    description: str,
    technical_specification: str,
    quantity: float,
    unit: str,
    status: str,
    basis: QuantityBasis,
    manufacturer: str = "",
    model: str = "",
    required_documents: Iterable[str] = (),
) -> ProcurementLine:
    override = _override_map(settings).get(item_id)
    final = float(override.quantity) if override is not None else float(quantity)
    rationale = override.rationale if override is not None else ""
    return ProcurementLine(
        item_id=item_id,
        category=category,
        description=description,
        technical_specification=technical_specification,
        auto_quantity=float(quantity),
        final_quantity=final,
        unit=unit,
        status=status,
        basis=basis,
        override_rationale=rationale,
        manufacturer=manufacturer,
        model=model,
        required_documents=tuple(required_documents),
    )


def _selected_svl_spec(project: ProjectData) -> tuple[str, str, str]:
    selected = project.svl.selected_candidate_id
    for candidate in project.svl.candidates:
        if candidate.candidate_id == selected:
            spec = (
                f"{candidate.technology}; MCOV {candidate.mcov_rms_v:g} V rms; "
                f"TOV 1 s {candidate.tov_1s_rms_v:g} V rms; residual {candidate.residual_voltage_peak_v:g} Vpk; "
                f"enerji {candidate.energy_capacity_kj:g} kJ; In {candidate.nominal_discharge_current_ka:g} kA"
            )
            return candidate.manufacturer, candidate.model, spec
    spec = (
        f"Bağlantı {project.svl.connection_mode}; sürekli gerilim marjı "
        f"%{project.svl.continuous_voltage_margin_percent:g}; nihai MCOV/TOV/residual/enerji sınıfı "
        "üretici eğrileri ve hesap sonucu ile seçilecektir."
    )
    return "", "", spec


def _build_drum_cuts(
    project: ProjectData,
    settings: ProcurementData,
    route_length_m: float,
    circuit_count: int,
    parallel: int,
) -> tuple[list[DrumCut], list[str]]:
    """Build only physically continuous installation cuts.

    Order allowance and operational spare are intentionally not represented as
    ``DrumCut`` objects.  They are procurement allocations, not route segments.
    """
    segments = _route_segments(project, route_length_m)
    cut_round = max(0.1, float(settings.drum_length_rounding_m))
    cuts: list[DrumCut] = []
    notes: list[str] = []
    for circuit in range(1, circuit_count + 1):
        for parallel_no in range(1, parallel + 1):
            for phase in PHASES:
                for segment_index, segment_length in enumerate(segments, start=1):
                    left_allowance = (
                        settings.termination_tail_m_per_end
                        if segment_index == 1
                        else settings.joint_tail_m_per_side
                    )
                    right_allowance = (
                        settings.termination_tail_m_per_end
                        if segment_index == len(segments)
                        else settings.joint_tail_m_per_side
                    )
                    end_allowance = float(left_allowance + right_allowance)
                    installation = float(segment_length) * settings.installation_allowance_percent / 100.0
                    required = _round_up(float(segment_length) + end_allowance + installation, cut_round)
                    cuts.append(DrumCut(
                        cut_id=f"CUT-C{circuit}-P{phase}-G{parallel_no}-S{segment_index}",
                        route_reference=" + ".join(section.name for section in _active_routes(project)),
                        circuit_no=circuit,
                        phase=phase,
                        segment_no=segment_index,
                        segment_length_m=float(segment_length),
                        end_allowance_m=end_allowance,
                        installation_allowance_m=installation,
                        required_cut_length_m=required,
                        notes="Terminasyon/joint kuyrukları ve montaj payı dahil.",
                    ))
    return cuts, notes


def _distribute_allowance(
    assignments: list[dict[str, Any]],
    amount_m: float,
    maximum_length_m: float,
    increment_m: float,
) -> float:
    """Distribute order allowance in physical drum-length increments."""
    remaining = max(0.0, float(amount_m))
    increment = max(1.0e-9, float(increment_m))
    while remaining >= increment - 1.0e-9:
        eligible = [item for item in assignments if maximum_length_m - item["loaded"] >= increment - 1.0e-9]
        if not eligible:
            break
        item = min(eligible, key=lambda candidate: (candidate["loaded"], candidate["route_cut"]))
        item["order_allowance"] += increment
        item["loaded"] += increment
        remaining -= increment
    if remaining > 1.0e-9:
        eligible = [item for item in assignments if maximum_length_m - item["loaded"] >= remaining - 1.0e-9]
        if eligible:
            item = min(eligible, key=lambda candidate: (candidate["loaded"], candidate["route_cut"]))
            item["order_allowance"] += remaining
            item["loaded"] += remaining
            remaining = 0.0
    return max(0.0, remaining)

def _pack_drums(
    cuts: list[DrumCut],
    maximum_length_m: float,
    mass_kg_km: float,
    cable_item_id: str,
    *,
    order_quantity_m: float,
    spare_stock_total_m: float,
    allocation_increment_m: float = 1.0,
) -> tuple[list[DrumAssignment], list[UnassignedDrumCut], DrumPlanSummary, list[str]]:
    warnings: list[str] = []
    tolerance = 1.0e-6
    route_cut_total = sum(item.required_cut_length_m for item in cuts)
    spare_target = max(0.0, min(float(spare_stock_total_m), max(0.0, order_quantity_m)))
    allowance_target = max(0.0, float(order_quantity_m) - route_cut_total - spare_target)

    unassigned: list[UnassignedDrumCut] = []
    normal: list[DrumCut] = []
    for item in cuts:
        if item.required_cut_length_m > maximum_length_m + tolerance:
            unassigned.append(UnassignedDrumCut(
                cut_id=item.cut_id,
                required_cut_length_m=item.required_cut_length_m,
                maximum_length_m=maximum_length_m,
                deficit_m=item.required_cut_length_m - maximum_length_m,
                reason="DRUM_CUT_EXCEEDS_MAXIMUM",
                route_reference=item.route_reference,
            ))
            warnings.append(
                f"{item.cut_id} için {item.required_cut_length_m:.1f} m kesim, izin verilen "
                f"{maximum_length_m:.1f} m makara boyunu aşıyor; fiziksel makara üretilmedi."
            )
        else:
            normal.append(item)

    bins: list[dict[str, Any]] = []
    for cut in sorted(normal, key=lambda item: item.required_cut_length_m, reverse=True):
        placed = False
        for item in bins:
            if item["loaded"] + cut.required_cut_length_m <= maximum_length_m + tolerance:
                item["cuts"].append(cut)
                item["route_cut"] += cut.required_cut_length_m
                item["loaded"] += cut.required_cut_length_m
                placed = True
                break
        if not placed:
            bins.append({
                "cuts": [cut],
                "route_cut": cut.required_cut_length_m,
                "order_allowance": 0.0,
                "spare": 0.0,
                "rounding": 0.0,
                "loaded": cut.required_cut_length_m,
                "kind": "INSTALLATION",
            })

    allowance_remaining = _distribute_allowance(
        bins, allowance_target, maximum_length_m, allocation_increment_m
    )
    while allowance_remaining > tolerance:
        addition = min(maximum_length_m, allowance_remaining)
        bins.append({
            "cuts": [], "route_cut": 0.0, "order_allowance": addition,
            "spare": 0.0, "rounding": 0.0, "loaded": addition,
            "kind": "ORDER_ALLOWANCE",
        })
        allowance_remaining -= addition

    spare_remaining = spare_target
    while spare_remaining > tolerance:
        addition = min(maximum_length_m, spare_remaining)
        bins.append({
            "cuts": [], "route_cut": 0.0, "order_allowance": 0.0,
            "spare": addition, "rounding": 0.0, "loaded": addition,
            "kind": "OPERATIONAL_SPARE",
        })
        spare_remaining -= addition

    assignments: list[DrumAssignment] = []
    for idx, item in enumerate(bins, start=1):
        load = float(item["loaded"])
        balance = maximum_length_m - load
        overload = max(0.0, -balance)
        remaining_capacity = max(0.0, balance)
        assignment_status = "INVALID_OVERLOAD" if overload > tolerance else "VALID"
        notes = (
            "Gerçek güzergâh kesimleri FFD ile yerleştirildi; sipariş/fire payı mevcut kapasitelere "
            "dağıtıldı; işletme yedeği ayrı stok makaralarında tutuldu."
        )
        assignments.append(DrumAssignment(
            drum_id=f"DRM-{idx:03d}",
            cable_item_id=cable_item_id,
            maximum_length_m=maximum_length_m,
            loaded_length_m=load,
            remaining_length_m=remaining_capacity,
            cuts=tuple(item["cuts"]),
            route_cut_length_m=float(item["route_cut"]),
            order_allowance_m=float(item["order_allowance"]),
            spare_stock_length_m=float(item["spare"]),
            rounding_reconciliation_m=float(item["rounding"]),
            capacity_balance_m=balance,
            remaining_capacity_m=remaining_capacity,
            overload_m=overload,
            estimated_net_mass_kg=load / 1000.0 * mass_kg_km if mass_kg_km > 0 else 0.0,
            assignment_status=assignment_status,
            status=STATUS_CONDITIONAL if assignment_status != "VALID" else STATUS_CONFIRMED,
            notes=notes,
        ))

    allocated = sum(item.loaded_length_m for item in assignments)
    unallocated = max(0.0, float(order_quantity_m) - allocated)
    overload_total = sum(item.overload_m for item in assignments)
    accounting_residual = allocated + unallocated - float(order_quantity_m)
    invalid_count = sum(item.assignment_status != "VALID" for item in assignments)
    if unassigned or invalid_count or abs(accounting_residual) > tolerance:
        plan_status = "INVALID"
    elif unallocated > tolerance:
        plan_status = "INCOMPLETE"
    else:
        plan_status = "VALID"
    summary = DrumPlanSummary(
        route_cut_total_m=route_cut_total,
        order_allowance_total_m=sum(item.order_allowance_m for item in assignments),
        spare_stock_total_m=sum(item.spare_stock_length_m for item in assignments),
        rounding_reconciliation_total_m=sum(item.rounding_reconciliation_m for item in assignments),
        allocated_total_m=allocated,
        unallocated_total_m=unallocated,
        overload_total_m=overload_total,
        maximum_single_drum_overload_m=max((item.overload_m for item in assignments), default=0.0),
        accounting_residual_m=accounting_residual,
        valid_drum_count=len(assignments) - invalid_count,
        invalid_drum_count=invalid_count,
        unassigned_cut_count=len(unassigned),
        drum_plan_status=plan_status,
    )
    return assignments, unassigned, summary, warnings

def _thermal_civil_lines(project: ProjectData, settings: ProcurementData) -> tuple[list[ProcurementLine], list[str]]:
    if not settings.include_civil_items:
        return [], []
    templates = {item.template_id: item for item in project.thermal_design.templates}
    lines: list[ProcurementLine] = []
    warnings: list[str] = []
    excavation = 0.0
    backfill = 0.0
    duct_length = 0.0
    route_cable_count = max(1, project.design_basis.circuit_count) * 3 * max(1, project.cable.parallel_cables_per_phase)
    used_regions = [region for region in project.thermal_design.regions if region.enabled]
    for region in used_regions:
        template = templates.get(region.template_id)
        if template is None:
            warnings.append(f"{region.name}: termal kesit şablonu bulunamadığı için inşaat metrajı üretilemedi.")
            continue
        length = float(region.length_m)
        width = float(region.overrides.get("trench_width_m", template.trench_width_m))
        depth = float(region.overrides.get("trench_depth_m", template.trench_depth_m))
        bedding = float(region.overrides.get("bedding_thickness_m", template.bedding_thickness_m))
        cover = float(region.overrides.get("cable_cover_height_m", template.cable_cover_height_m))
        excavation += max(0.0, width * depth * length)
        backfill += max(0.0, width * (bedding + cover) * length)
        if template.installation_type in {THERMAL_INSTALL_DUCT_BANK, THERMAL_INSTALL_HDD}:
            duct_length += length * route_cable_count
    if excavation > 0:
        lines.append(_line(
            settings, "CIV-EXC-001", CATEGORY_CIVIL, "Kablo hendeği kazısı",
            "Termal bölge kesit genişliği × hendek derinliği × bölge uzunluğu; şev, kabarma ve nakliye katsayıları hariç.",
            excavation, "m³", STATUS_ASSUMPTION,
            QuantityBasis("ThermalDesignData.regions/templates", "Termal güzergâh", "Σ(trench_width × trench_depth × region_length)", STATUS_ASSUMPTION,
                          notes="Kesitler DESIGN seviyesindeyse teklif ön metrajıdır."),
        ))
    if backfill > 0:
        lines.append(_line(
            settings, "CIV-TBF-001", CATEGORY_CIVIL, "Termal dolgu malzemesi",
            "Bedding + kablo üstü termal dolgu hacmi; sıkışma/kabarma ve satın alma yoğunluğu ayrıca teklif şartında belirtilecektir.",
            backfill, "m³", STATUS_ASSUMPTION,
            QuantityBasis("ThermalDesignData.regions/templates", "Termal güzergâh", "Σ(width × (bedding + cover) × length)", STATUS_ASSUMPTION),
        ))
    if duct_length > 0:
        lines.append(_line(
            settings, "CIV-DUCT-001", CATEGORY_CIVIL, "Kablo koruma borusu / duct",
            "İç/dış çap ve malzeme ilgili termal kesit şablonundan; her tek damarlı kablo için ayrı duct.",
            duct_length, "m", STATUS_CONDITIONAL,
            QuantityBasis("ThermalCrossSectionTemplate", "Duct/HDD bölgeleri", "region_length × phases × circuits × cables/phase", STATUS_CONDITIONAL),
        ))
    return lines, warnings


def build_procurement_package(project: ProjectData, settings: ProcurementData | None = None) -> ProcurementPackage:
    """Derive a traceable BOQ/BOM/RFQ package from the current project state."""

    project = project_with_synchronized_installation_geometry(project)
    settings = settings or project.procurement
    active_routes = _active_routes(project)
    route_length = sum(max(0.0, float(item.length_m)) for item in active_routes)
    if route_length <= 0:
        raise ProcurementInputError("BOQ/BOM için pozitif yeraltı güzergâh uzunluğu gereklidir.")
    circuits = max(1, int(project.design_basis.circuit_count or 1))
    parallel = max(1, int(project.cable.parallel_cables_per_phase or 1))
    single_core_count = 3 * circuits * parallel
    base_cable = route_length * single_core_count

    termination_nodes = [node for node in project.bonding.nodes if node.node_type == "TERMINATION"]
    joint_nodes = [node for node in project.bonding.nodes if node.node_type == "SECTIONALIZING_JOINT"]
    termination_points = len(termination_nodes) if termination_nodes else 2
    termination_units_auto = termination_points * single_core_count
    joint_units_auto = len(joint_nodes) * single_core_count
    accessory_plan = resolve_bonding_accessory_plan(project.bonding)
    cross_link_box_units_auto = accessory_plan.cross_link_box_units_per_circuit * circuits
    grounding_link_box_units_auto = accessory_plan.grounding_link_box_units_per_circuit * circuits
    custom_link_box_units_auto = accessory_plan.custom_link_box_units_per_circuit * circuits
    link_box_units_auto = accessory_plan.total_link_box_units_per_circuit * circuits
    svl_set_units_auto = accessory_plan.svl_set_units_per_circuit * circuits
    svl_pole_units_auto = accessory_plan.svl_pole_units_per_circuit * circuits

    termination_tail = termination_units_auto * max(0.0, settings.termination_tail_m_per_end)
    joint_tail = joint_units_auto * 2.0 * max(0.0, settings.joint_tail_m_per_side)
    installation_allowance = base_cable * max(0.0, settings.installation_allowance_percent) / 100.0
    subtotal = base_cable + termination_tail + joint_tail + installation_allowance
    waste = subtotal * max(0.0, settings.waste_percent) / 100.0
    spare = base_cable * max(0.0, settings.spare_cable_percent) / 100.0
    order_cable = _round_up(subtotal + waste + spare, max(0.1, settings.drum_length_rounding_m))

    cable = project.cable
    cable_spec = (
        f"{cable.voltage_class}; {cable.construction_type}; {cable.conductor_material} "
        f"{cable.conductor_area_mm2:g} mm²; metalik ekran/kılıf {cable.sheath_cross_section_mm2:g} mm²; "
        f"{cable.insulation}; dış çap {cable.overall_diameter_mm:g} mm; Rdc20 {cable.dc_resistance_20_ohm_km:g} Ω/km"
    )
    route_ref = " + ".join(item.name for item in active_routes)
    data_status = STATUS_CONFIRMED if cable.data_status == "VERIFIED" else STATUS_CONDITIONAL
    lines: list[ProcurementLine] = []
    lines.append(_line(
        settings, "CBL-001", CATEGORY_CABLE, "Tek damarlı OG güç kablosu", cable_spec,
        order_cable, "m", data_status,
        QuantityBasis(
            "ProjectData.route_sections + CableData + BondingSystemData",
            route_ref,
            "route × 3 faz × devre × paralel + termination/joint kuyrukları + montaj + fire + yedek",
            data_status,
            source_reference=cable.snapshot_hash or cable.catalog_record_id,
            notes=(
                f"Net tek-damarlı uzunluk {base_cable:.1f} m; termination kuyruğu {termination_tail:.1f} m; "
                f"joint kuyruğu {joint_tail:.1f} m; montaj payı {installation_allowance:.1f} m; "
                f"fire {waste:.1f} m; yedek {spare:.1f} m."
            ),
        ),
        manufacturer=cable.manufacturer,
        model=cable.model,
        required_documents=(
            "Üretici teknik veri föyü ve konstrüksiyon çizimi",
            "Rutin test sertifikaları",
            "Tip test raporları / standart uygunluk beyanı",
            "Makara boyu, makara ölçüsü ve net/brüt ağırlık",
            "Metalik ekran tel yapısı ve kısa devre dayanımı",
        ),
    ))

    termination_units = termination_units_auto + max(0, int(settings.spare_termination_units))
    lines.append(_line(
        settings, "ACC-TERM-001", CATEGORY_ACCESSORY, "Tek damarlı kablo terminasyonu",
        f"{cable.voltage_class}; {cable.conductor_area_mm2:g} mm² {cable.conductor_material}; iç/dış ortam ve tesis arayüzü proje uçlarına göre teyit edilecektir.",
        termination_units, "adet", STATUS_CONDITIONAL,
        QuantityBasis("BondingSystemData.nodes", route_ref, "TERMINATION nodes × 3 faz × devre × paralel + yedek", STATUS_CONDITIONAL),
        required_documents=("Tip test raporu", "Montaj talimatı", "Kablo uyumluluk çizelgesi", "Saha test gereksinimleri"),
    ))

    joint_units = joint_units_auto + max(0, int(settings.spare_joint_units))
    if joint_units > 0:
        lines.append(_line(
            settings, "ACC-JOINT-001", CATEGORY_ACCESSORY, "Tek damarlı sectionalizing / straight joint",
            f"{cable.voltage_class}; {cable.conductor_area_mm2:g} mm²; metalik ekran/kılıf ayrımı ve bonding çıkışları proje şemasına uygun.",
            joint_units, "adet", STATUS_CONDITIONAL,
            QuantityBasis("BondingSystemData.nodes", route_ref, "SECTIONALIZING_JOINT nodes × 3 faz × devre × paralel + yedek", STATUS_CONDITIONAL),
            required_documents=("Tip test raporu", "Montaj talimatı", "Ekran ayırma ve bonding bağlantı çizimi", "Joint bay minimum ölçüleri"),
        ))

    spare_link_boxes = max(0, int(settings.spare_link_box_units))
    link_box_units = link_box_units_auto + spare_link_boxes
    accessory_status = STATUS_CONDITIONAL if accessory_plan.status != ACCESSORY_VALID else STATUS_CONFIRMED
    if cross_link_box_units_auto > 0:
        lines.append(_line(
            settings, "BND-LB-CROSS-001", CATEGORY_BONDING, "Cross-bonding link box",
            f"{project.bonding.scheme}; üç faz çapraz bağlantı baraları ve SVL bağlantı noktaları; erişilebilirlik/IP sınıfı saha koşuluna göre.",
            cross_link_box_units_auto, "adet", accessory_status,
            QuantityBasis("BondingAccessoryPlan", route_ref, "minor cross boundaries × circuit count", accessory_status),
            required_documents=("İç çapraz bağlantı şeması", "Muhafaza/IP ve korozyon sınıfı", "Kısa devre akım dayanımı", "Kablo gland/terminal detayları"),
        ))
    if grounding_link_box_units_auto > 0:
        lines.append(_line(
            settings, "BND-LB-GROUND-001", CATEGORY_BONDING, "Major-section grounding link box",
            "Üç faz metalik kılıfın solid-ground bağlantısı; SVL içermez; çıkarılabilir test linkleri ve topraklama barası.",
            grounding_link_box_units_auto, "adet", accessory_status,
            QuantityBasis("BondingAccessoryPlan", route_ref, "major-section grounded boundaries × circuit count", accessory_status),
            required_documents=("Solid-ground iç bağlantı şeması", "Topraklama barası kısa devre dayanımı", "Muhafaza/IP ve korozyon sınıfı"),
        ))
    if spare_link_boxes > 0:
        lines.append(_line(
            settings, "BND-LB-SPARE-001", CATEGORY_BONDING, "Yedek bonding link box",
            "Link-box tipi satın alma öncesinde cross/grounding dağılımına göre açıkça seçilecektir.",
            spare_link_boxes, "adet", STATUS_ASSUMPTION,
            QuantityBasis("ProcurementData.spare_link_box_units", route_ref, "explicit spare link-box units", STATUS_ASSUMPTION),
        ))

    svl_set_units = svl_set_units_auto
    svl_pole_units = svl_pole_units_auto + max(0, int(settings.spare_svl_units))
    svl_units = svl_pole_units  # Backward-compatible meaning: independent SVL poles/elements.
    svl_manufacturer, svl_model, svl_spec = _selected_svl_spec(project)
    if svl_pole_units > 0:
        svl_status = STATUS_CONFIRMED if project.svl.selected_candidate_id and accessory_plan.status == ACCESSORY_VALID else STATUS_CONDITIONAL
        lines.append(_line(
            settings, "BND-SVL-001", CATEGORY_BONDING, "Metalik kılıf gerilim sınırlayıcı (SVL) elemanı/polü", svl_spec,
            svl_pole_units, "adet", svl_status,
            QuantityBasis("BondingAccessoryPlan + SvlSystemData", route_ref, "cross-bonding boundary × 3 pole × circuit + explicit spare poles", svl_status),
            manufacturer=svl_manufacturer,
            model=svl_model,
            required_documents=("MCOV/TOV eğrileri", "Residual gerilim eğrileri", "Enerji dayanım verileri", "Rutin test sertifikası", "Montaj ve topraklama talimatı"),
        ))
        lines.append(_line(
            settings, "BND-SVL-SET-001", CATEGORY_BONDING, "Üç fazlı SVL seti",
            "Bir cross-bonding link box için üç bağımsız SVL koruma kolu.",
            svl_set_units, "set", svl_status,
            QuantityBasis("BondingAccessoryPlan", route_ref, "cross-bonding link-box count", svl_status),
        ))

    lead_auto = sum(max(0.0, float(box.lead_length_m)) for box in project.bonding.link_boxes) * 3 * circuits
    lead_final_auto = _round_up(lead_auto * (1.0 + settings.bonding_lead_allowance_percent / 100.0), 1.0)
    if lead_final_auto > 0:
        lead_types = ", ".join(sorted({box.lead_type for box in project.bonding.link_boxes}))
        lines.append(_line(
            settings, "BND-LEAD-001", CATEGORY_BONDING, "Bonding bağlantı iletkeni / koaksiyel bonding kablosu",
            f"Tip: {lead_types or 'proje seçimi'}; olabildiğince kısa güzergâh; kesit, ekranlama, darbe ve kısa devre dayanımı üretici/hesap ile doğrulanacaktır.",
            lead_final_auto, "m", STATUS_CONDITIONAL,
            QuantityBasis("BondingLinkBox.lead_length_m", route_ref, "Σ(lead length × 3 faz × devre) + açık bonding lead payı", STATUS_CONDITIONAL),
        ))

    if settings.include_grounding_items:
        grounded_points = sum(1 for node in project.bonding.nodes if node.grounded) * circuits
        if grounded_points > 0:
            lines.append(_line(
                settings, "GND-POINT-001", CATEGORY_GROUNDING, "Bonding/terminasyon topraklama bağlantı noktası",
                "Topraklama iletkeni, bağlantı barası, pabuç ve korozyon koruması dahil; hedef direnç ve EPR kriteri topraklama tasarımından.",
                grounded_points, "nokta", STATUS_CONDITIONAL,
                QuantityBasis("BondingNode.grounded", route_ref, "grounded nodes × circuit count", STATUS_CONDITIONAL),
            ))
        if project.bonding.gcc_enabled:
            gcc_length = _round_up(route_length * circuits * (1.0 + settings.installation_allowance_percent / 100.0), 1.0)
            lines.append(_line(
                settings, "GND-GCC-001", CATEGORY_GROUNDING, "GCC/ECC toprak süreklilik iletkeni",
                f"{project.bonding.gcc_material} {project.bonding.gcc_area_mm2:g} mm²; işletme sıcaklığı {project.bonding.gcc_operating_temperature_c:g} °C.",
                gcc_length, "m", STATUS_CONDITIONAL,
                QuantityBasis("BondingSystemData.gcc_*", route_ref, "route × circuits + installation allowance", STATUS_CONDITIONAL),
            ))

    if settings.include_marking_accessories:
        spacing = max(0.1, float(settings.cable_cleat_spacing_m))
        cleat_sets = sum((math.ceil(max(0.0, section.length_m) / spacing) + 1) for section in active_routes) * circuits * parallel
        arrangement = project.cable.arrangement or project.design_basis.installation_profile
        lines.append(_line(
            settings, "SUP-CLEAT-001", CATEGORY_MARKING, "Üç faz kablo kelepçe/cleat seti",
            f"{arrangement}; azami aralık {spacing:g} m; kısa devre elektrodinamik dayanımı arıza hesabına göre belgelenecektir.",
            cleat_sets, "set", STATUS_CONDITIONAL,
            QuantityBasis("RouteSection.length_m + ProcurementData.cable_cleat_spacing_m", route_ref, "Σ(ceil(section/spacing)+1) × circuits × parallel", STATUS_CONDITIONAL),
            required_documents=("Kısa devre dinamik test/hesap raporu", "Kablo dış çapı uyumluluğu", "Montaj aralığı ve tork talimatı"),
        ))
        tape = _round_up(route_length * (1.0 + settings.warning_tape_allowance_percent / 100.0), 1.0)
        lines.append(_line(
            settings, "MRK-TAPE-001", CATEGORY_MARKING, "Yeraltı kablo ikaz bandı",
            "Güzergâh boyunca sürekli, dayanıklı ve proje dilinde işaretli; yatay bant sayısı tip kesite göre teyit edilecektir.",
            tape, "m", STATUS_ASSUMPTION,
            QuantityBasis("RouteSection.length_m", route_ref, "route length + warning tape allowance", STATUS_ASSUMPTION),
        ))
        marker_count = max(2, len(active_routes) + len(project.bonding.nodes))
        lines.append(_line(
            settings, "MRK-TAG-001", CATEGORY_MARKING, "Kablo güzergâh ve aksesuar işaretleme seti",
            "Terminasyon, joint, link box ve güzergâh dönüş/erişim noktaları için UV/korozyon dayanımlı etiketleme.",
            marker_count, "set", STATUS_ASSUMPTION,
            QuantityBasis("Route/Bonding graph", route_ref, "route sections + accessory points", STATUS_ASSUMPTION),
        ))

    civil_lines, civil_warnings = _thermal_civil_lines(project, settings)
    lines.extend(civil_lines)

    maximum_drum, mass_kg_km, drum_source = _catalog_delivery_and_mass(project, settings)
    cuts, cut_warnings = _build_drum_cuts(project, settings, route_length, circuits, parallel)
    spare_stock_target = _round_up(spare, max(0.1, settings.drum_length_rounding_m)) if spare > 0 else 0.0
    drums, unassigned_cuts, drum_plan, drum_warnings = _pack_drums(
        cuts,
        maximum_drum,
        mass_kg_km,
        "CBL-001",
        order_quantity_m=order_cable,
        spare_stock_total_m=spare_stock_target,
        allocation_increment_m=max(0.1, settings.drum_length_rounding_m),
    )

    warnings: list[str] = []
    warnings.extend(civil_warnings)
    warnings.extend(cut_warnings)
    warnings.extend(drum_warnings)
    warnings.extend(accessory_plan.errors)
    warnings.extend(accessory_plan.warnings)
    if drum_plan.drum_plan_status != "VALID":
        warnings.append(f"Makara planı {drum_plan.drum_plan_status}: aşım/atanmamış kesim/muhasebe kapıları kontrol edilmelidir.")
    if accessory_plan.status != ACCESSORY_VALID:
        warnings.append(f"Bonding aksesuar planı {accessory_plan.status}; BOQ link-box/SVL miktarları nihai kabul edilemez.")
    if cable.data_status != "VERIFIED":
        warnings.append("Kablo verisi VERIFIED değildir; BOQ/BOM/RFQ teknik tanımı koşullu proje verisi içerir.")
    if not project.svl.selected_candidate_id and svl_units > 0:
        warnings.append("SVL adedi bonding grafiğinden üretildi; üretici/model ve MCOV/TOV/enerji sınıfı nihai seçilmemiştir.")
    if project.thermal_design.active_data_state != "AS_BUILT":
        warnings.append("Kazı ve termal dolgu miktarları DESIGN termal kesitlerinden türetilmiştir; saha/as-built metrajı değildir.")
    if settings.quantity_overrides:
        warnings.append("Bir veya daha fazla otomatik miktar kullanıcı tarafından gerekçeli olarak değiştirilmiştir; RFQ'da otomatik ve nihai miktarlar birlikte gösterilir.")

    assumptions = (
        f"Devre sayısı: {circuits}; paralel kablo/faz: {parallel}; tek damarlı iletken yolu sayısı: {single_core_count}.",
        f"Montaj payı %{settings.installation_allowance_percent:g}; fire %{settings.waste_percent:g}; yedek kablo %{settings.spare_cable_percent:g}.",
        f"Terminasyon kuyruğu {settings.termination_tail_m_per_end:g} m/uç; joint kuyruğu {settings.joint_tail_m_per_side:g} m/taraf.",
        f"Makara azami boyu {maximum_drum:g} m ({drum_source}); kesim yuvarlaması {settings.drum_length_rounding_m:g} m.",
        "Joint/termination miktarları tek damarlı aksesuar adedidir; link box miktarı üç fazlı kutu adedidir; SVL miktarı tek fazlı eleman adedidir.",
        "BOQ/BOM miktarları maliyet veya satın alma onayı değildir; teknik teklif, üretici uyumluluğu ve saha metrajı ayrıca kontrol edilmelidir.",
    )
    package_status = STATUS_CONDITIONAL if warnings else STATUS_CONFIRMED
    summary = ProcurementSummary(
        net_route_length_m=route_length,
        installed_single_core_length_m=subtotal,
        order_single_core_length_m=order_cable,
        circuit_count=circuits,
        parallel_cables_per_phase=parallel,
        termination_units=termination_units,
        joint_units=joint_units,
        link_box_units=link_box_units,
        svl_units=svl_units,
        drum_count=len(drums),
        status=package_status,
        cross_bonding_link_box_units=cross_link_box_units_auto,
        grounding_link_box_units=grounding_link_box_units_auto,
        custom_link_box_units=custom_link_box_units_auto,
        svl_set_units=svl_set_units,
        svl_pole_units=svl_pole_units,
        accessory_plan_status=accessory_plan.status,
        drum_plan_status=drum_plan.drum_plan_status,
        route_cut_total_m=drum_plan.route_cut_total_m,
        order_allowance_total_m=drum_plan.order_allowance_total_m,
        spare_stock_total_m=drum_plan.spare_stock_total_m,
        allocated_total_m=drum_plan.allocated_total_m,
        unallocated_total_m=drum_plan.unallocated_total_m,
        overload_total_m=drum_plan.overload_total_m,
        accounting_residual_m=drum_plan.accounting_residual_m,
        invalid_drum_count=drum_plan.invalid_drum_count,
        unassigned_cut_count=drum_plan.unassigned_cut_count,
    )
    return ProcurementPackage(
        reference=REFERENCE,
        generated_at=datetime.now().isoformat(timespec="seconds"),
        project_name=project.project_name,
        project_code=project.project_code,
        project_signature_sha256=_stable_project_signature(project),
        settings=settings,
        summary=summary,
        lines=tuple(lines),
        drums=tuple(drums),
        unassigned_route_cuts=tuple(unassigned_cuts),
        drum_plan=drum_plan,
        accessory_plan=accessory_plan,
        warnings=tuple(warnings),
        assumptions=assumptions,
    )


def render_procurement_markdown(package: ProcurementPackage) -> str:
    s = package.summary
    lines = [
        f"# {package.project_name} - BOQ/BOM/RFQ Paketi",
        "",
        f"- Proje kodu: `{package.project_code}`",
        f"- Üretim: `{package.generated_at}`",
        f"- Proje imzası: `{package.project_signature_sha256}`",
        f"- Durum: **{s.status}**",
        "",
        "## Özet",
        "",
        "| Gösterge | Değer |",
        "|---|---:|",
        f"| Net yeraltı güzergâhı | {s.net_route_length_m:.3f} m |",
        f"| Montajlı tek damarlı kablo | {s.installed_single_core_length_m:.3f} m |",
        f"| Sipariş tek damarlı kablo | {s.order_single_core_length_m:.3f} m |",
        f"| Termination | {s.termination_units} adet |",
        f"| Joint | {s.joint_units} adet |",
        f"| Cross-bonding link box | {s.cross_bonding_link_box_units} adet |",
        f"| Grounding link box | {s.grounding_link_box_units} adet |",
        f"| Toplam link box | {s.link_box_units} adet |",
        f"| SVL seti | {s.svl_set_units} set |",
        f"| SVL elemanı/polü | {s.svl_pole_units} adet |",
        f"| Makara | {s.drum_count} adet |",
        f"| Makara planı | {s.drum_plan_status} |",
        f"| Sipariş/fire payı | {s.order_allowance_total_m:.3f} m |",
        f"| Toplam aşım | {s.overload_total_m:.3f} m |",
        "",
        "## BOQ / BOM",
        "",
        "| Kalem | Kategori | Tanım | Otomatik miktar | Nihai miktar | Birim | Durum | Dayanak |",
        "|---|---|---|---:|---:|---|---|---|",
    ]
    for item in package.lines:
        lines.append(
            "| " + " | ".join([
                item.item_id,
                item.category,
                item.description.replace("|", "\\|"),
                _fmt(item.auto_quantity),
                _fmt(item.final_quantity),
                item.unit,
                item.status,
                item.basis.formula.replace("|", "\\|"),
            ]) + " |"
        )
    lines.extend(["", "## RFQ Teknik Gereksinimleri", ""])
    for item in package.lines_for_view(VIEW_RFQ):
        lines.extend([
            f"### {item.item_id} - {item.description}",
            "",
            item.technical_specification,
            "",
            f"Miktar: **{_fmt(item.final_quantity)} {item.unit}**",
            "",
        ])
        if item.required_documents:
            lines.append("İstenen belgeler:")
            lines.extend(f"- {doc}" for doc in item.required_documents)
            lines.append("")
    lines.extend([
        "## Makara Planı",
        "",
        "| Makara | Kesim | Fire/pay | Yedek | Toplam | Azami | Bakiye | Kalan | Aşım | Kesimler | Durum |",
        "|---|---:|---:|---:|---:|---:|---:|---:|---:|---|---|",
    ])
    for drum in package.drums:
        cut_text = ", ".join(f"{cut.cut_id}:{cut.required_cut_length_m:g}m" for cut in drum.cuts)
        lines.append(
            f"| {drum.drum_id} | {drum.route_cut_length_m:.1f} m | {drum.order_allowance_m:.1f} m | "
            f"{drum.spare_stock_length_m:.1f} m | {drum.loaded_length_m:.1f} m | {drum.maximum_length_m:.1f} m | "
            f"{drum.capacity_balance_m:.1f} m | {drum.remaining_capacity_m:.1f} m | {drum.overload_m:.1f} m | "
            f"{cut_text} | {drum.assignment_status} |"
        )
    lines.extend(["", "## Varsayımlar", ""])
    lines.extend(f"- {item}" for item in package.assumptions)
    lines.extend(["", "## Uyarılar", ""])
    lines.extend(f"- {item}" for item in package.warnings or ("Açık uyarı yok.",))
    return "\n".join(lines).rstrip() + "\n"


def render_procurement_html(package: ProcurementPackage) -> str:
    def rows(items: Iterable[ProcurementLine]) -> str:
        return "".join(
            "<tr>" + "".join(f"<td>{escape(str(value))}</td>" for value in (
                item.item_id, item.category, item.description, _fmt(item.auto_quantity),
                _fmt(item.final_quantity), item.unit, item.status, item.basis.formula,
            )) + "</tr>" for item in items
        )
    drum_rows = "".join(
        "<tr>" + "".join(f"<td>{escape(str(value))}</td>" for value in (
            drum.drum_id,
            _fmt(drum.route_cut_length_m, 1),
            _fmt(drum.order_allowance_m, 1),
            _fmt(drum.spare_stock_length_m, 1),
            _fmt(drum.loaded_length_m, 1),
            _fmt(drum.maximum_length_m, 1),
            _fmt(drum.capacity_balance_m, 1),
            _fmt(drum.remaining_capacity_m, 1),
            _fmt(drum.overload_m, 1),
            ", ".join(f"{cut.cut_id}: {cut.required_cut_length_m:g} m" for cut in drum.cuts),
            drum.assignment_status,
        )) + "</tr>" for drum in package.drums
    )
    warning_html = "".join(f"<li>{escape(item)}</li>" for item in package.warnings)
    assumption_html = "".join(f"<li>{escape(item)}</li>" for item in package.assumptions)
    s = package.summary
    return f"""<!doctype html><html lang='tr'><head><meta charset='utf-8'><title>{escape(package.project_code)} BOQ BOM RFQ</title>
<style>
@page {{ size:A4 landscape; margin:12mm; }}
:root{{--navy:#17324a;--navy2:#294e69;--line:#cbd5de;--soft:#eef3f7;--text:#1f2933;--warning:#fff2cf;}}
*{{box-sizing:border-box}} body{{font-family:Arial,'Noto Sans',sans-serif;color:var(--text);margin:0;background:#f4f6f8}}
.report{{max-width:1400px;margin:20px auto;background:#fff;padding:28px;box-shadow:0 4px 18px #0002}}
.banner{{background:var(--navy);color:#fff;padding:18px 22px}} .banner h1{{margin:0;color:#fff;font-size:25px}} .banner p{{margin:5px 0 0;color:#fff}}
h2{{background:var(--navy);color:#fff;padding:8px 12px;font-size:17px;margin-top:28px}} h3{{color:var(--navy2)}}
table{{width:100%;border-collapse:collapse;font-size:11px}} th{{background:var(--navy);color:#fff;text-align:left}} th,td{{border:1px solid var(--line);padding:5px 6px;vertical-align:top}} tbody tr:nth-child(even){{background:#f8fafb}}
.kpis{{display:grid;grid-template-columns:repeat(4,1fr);gap:8px;margin:16px 0}} .kpi{{border:1px solid var(--line);padding:9px;background:var(--soft)}} .kpi b{{display:block;color:var(--navy)}}
.warning{{background:var(--warning);border-left:5px solid #c98900;padding:8px 12px}} code{{word-break:break-all}}
@media print{{body{{background:#fff}}.report{{margin:0;padding:0;box-shadow:none}}tr{{page-break-inside:avoid}}}}
</style></head><body><main class='report'><div class='banner'><h1>DiTuS Kablo Analizör™ - BOQ/BOM/RFQ</h1><p>{escape(package.project_name)} · {escape(package.project_code)} · {escape(s.status)}</p></div>
<div class='kpis'><div class='kpi'><b>Net güzergâh</b>{s.net_route_length_m:.1f} m</div><div class='kpi'><b>Sipariş kablo</b>{s.order_single_core_length_m:.1f} m</div><div class='kpi'><b>Makara</b>{s.drum_count}</div><div class='kpi'><b>Proje imzası</b><code>{escape(package.project_signature_sha256[:16])}…</code></div></div>
<h2>BOQ / BOM</h2><table><thead><tr><th>Kalem</th><th>Kategori</th><th>Tanım</th><th>Otomatik</th><th>Nihai</th><th>Birim</th><th>Durum</th><th>Dayanak</th></tr></thead><tbody>{rows(package.lines)}</tbody></table>
<h2>Makara Planı</h2><table><thead><tr><th>Makara</th><th>Kesim [m]</th><th>Fire/pay [m]</th><th>Yedek [m]</th><th>Toplam [m]</th><th>Azami [m]</th><th>Bakiye [m]</th><th>Kalan [m]</th><th>Aşım [m]</th><th>Kesimler</th><th>Durum</th></tr></thead><tbody>{drum_rows}</tbody></table>
<h2>Varsayımlar</h2><ul>{assumption_html}</ul><h2>Uyarılar</h2><div class='warning'><ul>{warning_html or '<li>Açık uyarı yok.</li>'}</ul></div>
<p><small>Üretim: {escape(package.generated_at)} · {escape(package.reference)}</small></p></main></body></html>"""


def _write_csv_bundle(package: ProcurementPackage, directory: Path, stem: str) -> dict[str, Path]:
    paths: dict[str, Path] = {}
    tables = {
        "boq": package.lines_for_view(VIEW_BOQ),
        "bom": package.lines_for_view(VIEW_BOM),
        "rfq": package.lines_for_view(VIEW_RFQ),
    }
    for name, items in tables.items():
        path = directory / f"{stem}_{name}.csv"
        with path.open("w", encoding="utf-8-sig", newline="") as handle:
            writer = csv.writer(handle, delimiter=";")
            writer.writerow(["Kalem No", "Kategori", "Tanım", "Teknik Özellik", "Otomatik Miktar", "Nihai Miktar", "Birim", "Durum", "Kaynak Nesne", "Güzergâh", "Formül", "Override Gerekçesi"])
            for item in items:
                writer.writerow([
                    item.item_id, item.category, item.description, item.technical_specification,
                    item.auto_quantity, item.final_quantity, item.unit, item.status,
                    item.basis.source_object, item.basis.route_reference, item.basis.formula, item.override_rationale,
                ])
        paths[name] = path
    drum_path = directory / f"{stem}_drum_plan.csv"
    with drum_path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.writer(handle, delimiter=";")
        writer.writerow([
            "Makara", "Azami m", "Güzergâh Kesimi m", "Sipariş/Fire Payı m", "İşletme Yedeği m",
            "Yuvarlama Mutabakatı m", "Yüklenen m", "Kapasite Bakiyesi m", "Kalan Kapasite m",
            "Aşım m", "Tahmini Net Ağırlık kg", "Kesimler", "Fiziksel Durum", "Veri Durumu", "Not"
        ])
        for drum in package.drums:
            writer.writerow([
                drum.drum_id, drum.maximum_length_m, drum.route_cut_length_m, drum.order_allowance_m,
                drum.spare_stock_length_m, drum.rounding_reconciliation_m, drum.loaded_length_m,
                drum.capacity_balance_m, drum.remaining_capacity_m, drum.overload_m,
                drum.estimated_net_mass_kg, ", ".join(f"{c.cut_id}:{c.required_cut_length_m:g}m" for c in drum.cuts),
                drum.assignment_status, drum.status, drum.notes,
            ])
    paths["drum_plan"] = drum_path
    return paths


def _write_xlsx(package: ProcurementPackage, path: Path) -> None:
    try:
        import xlsxwriter
    except ImportError as exc:
        raise ProcurementInputError("XLSX çıktısı için XlsxWriter kurulmalıdır.") from exc

    workbook = xlsxwriter.Workbook(path)
    workbook.set_properties({
        "title": f"{package.project_code} BOQ BOM RFQ",
        "subject": package.project_name,
        "comments": f"Generated by DiTuS Kablo Analizör v{__version__} procurement engine",
    })
    navy = "#17324A"
    white = "#FFFFFF"
    soft = "#EEF3F7"
    line = "#CBD5DE"
    warning = "#FFF2CF"
    title_fmt = workbook.add_format({"bold": True, "font_size": 18, "font_color": white, "bg_color": navy, "align": "left", "valign": "vcenter"})
    header_fmt = workbook.add_format({"bold": True, "font_color": white, "bg_color": navy, "border": 1, "border_color": line, "text_wrap": True, "valign": "top"})
    cell_fmt = workbook.add_format({"border": 1, "border_color": line, "valign": "top", "text_wrap": True})
    num_fmt = workbook.add_format({"border": 1, "border_color": line, "num_format": "0.000", "valign": "top"})
    money_fmt = workbook.add_format({"border": 1, "border_color": line, "num_format": "#,##0.00", "valign": "top"})
    label_fmt = workbook.add_format({"bold": True, "bg_color": soft, "border": 1, "border_color": line})
    warn_fmt = workbook.add_format({"bg_color": warning, "text_wrap": True, "valign": "top"})

    summary = workbook.add_worksheet("Özet")
    summary.set_column("A:A", 28)
    summary.set_column("B:B", 42)
    summary.merge_range("A1:F2", "DiTuS Kablo Analizör™ - BOQ/BOM/RFQ Paketi", title_fmt)
    summary.write("A4", "Proje", label_fmt); summary.write("B4", package.project_name, cell_fmt)
    summary.write("A5", "Proje kodu", label_fmt); summary.write("B5", package.project_code, cell_fmt)
    summary.write("A6", "Durum", label_fmt); summary.write("B6", package.summary.status, cell_fmt)
    summary.write("A7", "Proje SHA-256", label_fmt); summary.write("B7", package.project_signature_sha256, cell_fmt)
    summary.write("A8", "Üretim", label_fmt); summary.write("B8", package.generated_at, cell_fmt)
    metrics = [
        ("Net güzergâh [m]", package.summary.net_route_length_m),
        ("Montajlı tek damarlı kablo [m]", package.summary.installed_single_core_length_m),
        ("Sipariş tek damarlı kablo [m]", package.summary.order_single_core_length_m),
        ("Termination [adet]", package.summary.termination_units),
        ("Joint [adet]", package.summary.joint_units),
        ("Cross-bonding link box [adet]", package.summary.cross_bonding_link_box_units),
        ("Grounding link box [adet]", package.summary.grounding_link_box_units),
        ("Toplam link box [adet]", package.summary.link_box_units),
        ("SVL seti [set]", package.summary.svl_set_units),
        ("SVL elemanı/polü [adet]", package.summary.svl_pole_units),
        ("Makara [adet]", package.summary.drum_count),
        ("Makara planı durumu", package.summary.drum_plan_status),
        ("Sipariş/fire payı [m]", package.summary.order_allowance_total_m),
        ("Toplam aşım [m]", package.summary.overload_total_m),
    ]
    summary.write_row("A10", ["Gösterge", "Değer"], header_fmt)
    metric_start_row = 10  # zero-based; Excel row 11
    for row, (label, value) in enumerate(metrics, start=metric_start_row):
        summary.write(row, 0, label, cell_fmt)
        summary.write(row, 1, value, num_fmt if isinstance(value, (int, float)) else cell_fmt)
    section_header_row = metric_start_row + len(metrics) + 1
    section_data_row = section_header_row + 1
    summary.write(section_header_row, 0, "Varsayımlar", header_fmt)
    for row, text in enumerate(package.assumptions, start=section_data_row):
        summary.write(row, 0, text, cell_fmt)
    summary.write(section_header_row, 3, "Uyarılar", header_fmt)
    for row, text in enumerate(package.warnings or ("Açık uyarı yok.",), start=section_data_row):
        summary.write(row, 3, text, warn_fmt)
    summary.set_column("D:D", 68)
    summary.freeze_panes(3, 0)

    line_headers = ["Kalem No", "Kategori", "Tanım", "Teknik Özellik", "Otomatik Miktar", "Nihai Miktar", "Birim", "Durum", "Kaynak Nesne", "Güzergâh", "Formül/Dayanak", "Kaynak Referansı", "Override Gerekçesi"]
    for sheet_name, view in (("BOQ", VIEW_BOQ), ("BOM", VIEW_BOM)):
        ws = workbook.add_worksheet(sheet_name)
        ws.write_row(0, 0, line_headers, header_fmt)
        for row, item in enumerate(package.lines_for_view(view), start=1):
            values = [
                item.item_id, item.category, item.description, item.technical_specification,
                item.auto_quantity, item.final_quantity, item.unit, item.status,
                item.basis.source_object, item.basis.route_reference, item.basis.formula,
                item.basis.source_reference, item.override_rationale,
            ]
            for col, value in enumerate(values):
                ws.write(row, col, value, num_fmt if col in {4, 5} else cell_fmt)
        ws.autofilter(0, 0, max(1, len(package.lines_for_view(view))), len(line_headers)-1)
        ws.freeze_panes(1, 0)
        ws.set_column("A:B", 18); ws.set_column("C:C", 30); ws.set_column("D:D", 55)
        ws.set_column("E:F", 14); ws.set_column("G:H", 18); ws.set_column("I:M", 30)

    rfq = workbook.add_worksheet("RFQ")
    rfq_headers = line_headers + ["İstenen Belgeler", "Teklif Edilen Marka/Model", "Teknik Uygunluk", "Birim Fiyat", "Toplam Fiyat", "Teslim Süresi", "Tedarikçi Notu"]
    rfq.write_row(0, 0, rfq_headers, header_fmt)
    rfq_items = package.lines_for_view(VIEW_RFQ)
    for row, item in enumerate(rfq_items, start=1):
        fixed = [
            item.item_id, item.category, item.description, item.technical_specification,
            item.auto_quantity, item.final_quantity, item.unit, item.status,
            item.basis.source_object, item.basis.route_reference, item.basis.formula,
            item.basis.source_reference, item.override_rationale, "\n".join(item.required_documents),
        ]
        for col, value in enumerate(fixed): rfq.write(row, col, value, num_fmt if col in {4, 5} else cell_fmt)
        rfq.write_blank(row, 14, None, cell_fmt)
        rfq.write_blank(row, 15, None, cell_fmt)
        rfq.write_blank(row, 16, None, money_fmt)
        rfq.write_formula(row, 17, f"=F{row+1}*Q{row+1}", money_fmt)
        rfq.write_blank(row, 18, None, cell_fmt)
        rfq.write_blank(row, 19, None, cell_fmt)
    rfq.autofilter(0, 0, max(1, len(rfq_items)), len(rfq_headers)-1)
    rfq.freeze_panes(1, 0)
    rfq.set_column("A:B", 17); rfq.set_column("C:C", 30); rfq.set_column("D:D", 55)
    rfq.set_column("E:F", 14); rfq.set_column("G:H", 18); rfq.set_column("I:N", 28)
    rfq.set_column("O:P", 24); rfq.set_column("Q:R", 15); rfq.set_column("S:T", 22)

    drum = workbook.add_worksheet("Makara Planı")
    drum_headers = [
        "Makara", "Kablo Kalemi", "Azami Boy [m]", "Güzergâh Kesimi [m]", "Sipariş/Fire Payı [m]",
        "İşletme Yedeği [m]", "Yuvarlama [m]", "Yüklenen [m]", "Kapasite Bakiyesi [m]",
        "Kalan Kapasite [m]", "Aşım [m]", "Tahmini Net Ağırlık [kg]", "Kesimler",
        "Fiziksel Durum", "Veri Durumu", "Not"
    ]
    drum.write_row(0, 0, drum_headers, header_fmt)
    for row, item in enumerate(package.drums, start=1):
        values = [
            item.drum_id, item.cable_item_id, item.maximum_length_m, item.route_cut_length_m,
            item.order_allowance_m, item.spare_stock_length_m, item.rounding_reconciliation_m,
            item.loaded_length_m, item.capacity_balance_m, item.remaining_capacity_m, item.overload_m,
            item.estimated_net_mass_kg,
            "\n".join(f"{cut.cut_id}: {cut.required_cut_length_m:g} m" for cut in item.cuts),
            item.assignment_status, item.status, item.notes,
        ]
        for col, value in enumerate(values):
            drum.write(row, col, value, num_fmt if col in set(range(2, 12)) else cell_fmt)
    if package.drums:
        drum.conditional_format(1, 8, len(package.drums), 8, {"type": "cell", "criteria": "<", "value": 0, "format": warn_fmt})
        drum.conditional_format(1, 10, len(package.drums), 10, {"type": "cell", "criteria": ">", "value": 0, "format": warn_fmt})
    drum.freeze_panes(1, 0)
    drum.autofilter(0, 0, max(1, len(package.drums)), len(drum_headers)-1)
    drum.set_column("A:B", 16)
    drum.set_column("C:L", 17)
    drum.set_column("M:M", 45)
    drum.set_column("N:P", 24)

    cuts = workbook.add_worksheet("Kesim Listesi")
    cut_headers = ["Kesim", "Güzergâh", "Devre", "Faz", "Segment", "Net Segment [m]", "Uç Payı [m]", "Montaj Payı [m]", "Kesim Boyu [m]", "Tip", "Not"]
    cuts.write_row(0,0,cut_headers,header_fmt)
    all_cuts = [cut for item in package.drums for cut in item.cuts]
    for row, cut in enumerate(all_cuts, start=1):
        values = [cut.cut_id, cut.route_reference, cut.circuit_no, cut.phase, cut.segment_no, cut.segment_length_m, cut.end_allowance_m, cut.installation_allowance_m, cut.required_cut_length_m, cut.cut_type, cut.notes]
        for col, value in enumerate(values): cuts.write(row,col,value,num_fmt if col in {5,6,7,8} else cell_fmt)
    cuts.freeze_panes(1,0); cuts.autofilter(0,0,max(1,len(all_cuts)),len(cut_headers)-1)
    cuts.set_column("A:A",25); cuts.set_column("B:B",45); cuts.set_column("C:E",10); cuts.set_column("F:I",16); cuts.set_column("J:K",28)

    workbook.close()


def _write_docx(package: ProcurementPackage, path: Path) -> None:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt, RGBColor
    except ImportError as exc:
        raise ProcurementInputError("DOCX çıktısı için python-docx kurulmalıdır.") from exc

    navy = "17324A"
    white = RGBColor(255,255,255)
    doc = Document()
    section = doc.sections[0]
    section.orientation = 1
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.4); section.bottom_margin = Cm(1.4); section.left_margin = Cm(1.3); section.right_margin = Cm(1.3)
    doc.styles["Normal"].font.name = "Liberation Sans"; doc.styles["Normal"].font.size = Pt(8)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DiTuS Kablo Analizör™ - BOQ/BOM/RFQ Paketi"); r.bold = True; r.font.size = Pt(20)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run(f"{package.project_name} · {package.project_code}")

    def shade_cell(cell, fill: str, font_color: RGBColor | None = None, bold: bool = False) -> None:
        tc_pr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); tc_pr.append(shd)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = bold
                if font_color is not None: run.font.color.rgb = font_color

    def add_banner(text: str) -> None:
        table = doc.add_table(rows=1, cols=1); cell=table.cell(0,0); cell.text=text; shade_cell(cell, navy, white, True)

    add_banner("Özet")
    summary = doc.add_table(rows=0, cols=2); summary.style="Table Grid"
    for label, value in (
        ("Net güzergâh", f"{package.summary.net_route_length_m:.3f} m"),
        ("Sipariş tek damarlı kablo", f"{package.summary.order_single_core_length_m:.3f} m"),
        ("Termination / Joint", f"{package.summary.termination_units} / {package.summary.joint_units}"),
        ("Cross / Ground / Toplam Link Box", f"{package.summary.cross_bonding_link_box_units} / {package.summary.grounding_link_box_units} / {package.summary.link_box_units}"),
        ("SVL set / eleman", f"{package.summary.svl_set_units} / {package.summary.svl_pole_units}"),
        ("Makara / plan", f"{package.summary.drum_count} / {package.summary.drum_plan_status}"),
        ("Sipariş/fire payı / aşım", f"{package.summary.order_allowance_total_m:.1f} m / {package.summary.overload_total_m:.1f} m"),
        ("Durum", package.summary.status),
        ("Proje SHA-256", package.project_signature_sha256),
    ):
        c=summary.add_row().cells; c[0].text=label; c[1].text=value; shade_cell(c[0],"EEF3F7",None,True)

    add_banner("BOQ / BOM")
    headers=("Kalem","Kategori","Tanım","Teknik özellik","Otomatik","Nihai","Birim","Durum","Dayanak")
    table=doc.add_table(rows=1,cols=len(headers)); table.style="Table Grid"
    for i,h in enumerate(headers): table.rows[0].cells[i].text=h; shade_cell(table.rows[0].cells[i],navy,white,True)
    for item in package.lines:
        c=table.add_row().cells
        values=(item.item_id,item.category,item.description,item.technical_specification,_fmt(item.auto_quantity),_fmt(item.final_quantity),item.unit,item.status,item.basis.formula)
        for i,v in enumerate(values): c[i].text=str(v)

    add_banner("Makara Planı")
    headers=("Makara","Kesim","Fire/pay","Yedek","Toplam","Azami","Bakiye","Aşım","Kesimler","Durum")
    table=doc.add_table(rows=1,cols=len(headers)); table.style="Table Grid"
    for i,h in enumerate(headers): table.rows[0].cells[i].text=h; shade_cell(table.rows[0].cells[i],navy,white,True)
    for drum in package.drums:
        c=table.add_row().cells
        values=(drum.drum_id,f"{drum.route_cut_length_m:.1f}",f"{drum.order_allowance_m:.1f}",f"{drum.spare_stock_length_m:.1f}",f"{drum.loaded_length_m:.1f}",f"{drum.maximum_length_m:.1f}",f"{drum.capacity_balance_m:.1f}",f"{drum.overload_m:.1f}",", ".join(f"{x.cut_id}:{x.required_cut_length_m:g}m" for x in drum.cuts),drum.assignment_status)
        for i,v in enumerate(values): c[i].text=str(v)

    add_banner("Varsayımlar ve Uyarılar")
    for item in package.assumptions: doc.add_paragraph(item,style="List Bullet")
    for item in package.warnings: doc.add_paragraph(item,style="List Bullet")
    footer=section.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; footer.add_run(f"DiTuS Kablo Analizör™ · {package.project_code} · {package.generated_at}")
    doc.core_properties.title=f"{package.project_code} BOQ BOM RFQ"
    path.parent.mkdir(parents=True,exist_ok=True); doc.save(path)


def _write_pdf(package: ProcurementPackage, path: Path) -> None:
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    except ImportError as exc:
        raise ProcurementInputError("PDF çıktısı için reportlab kurulmalıdır.") from exc
    regular="/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"; bold="/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"
    if Path(regular).exists():
        pdfmetrics.registerFont(TTFont("DiTuS",regular)); pdfmetrics.registerFont(TTFont("DiTuS-Bold",bold)); normal,bold_name="DiTuS","DiTuS-Bold"
    else: normal,bold_name="Helvetica","Helvetica-Bold"
    navy=colors.HexColor("#17324A"); white=colors.white; line=colors.HexColor("#CBD5DE")
    doc=SimpleDocTemplate(str(path),pagesize=landscape(A4),leftMargin=10*mm,rightMargin=10*mm,topMargin=10*mm,bottomMargin=12*mm)
    styles=getSampleStyleSheet(); title=ParagraphStyle("T",parent=styles["Title"],fontName=bold_name,fontSize=20,textColor=navy); body=ParagraphStyle("B",parent=styles["BodyText"],fontName=normal,fontSize=6.6,leading=8); small=ParagraphStyle("S",parent=body,fontSize=5.8,leading=7); table_header=ParagraphStyle("TH",parent=small,fontName=bold_name,textColor=white); banner=ParagraphStyle("H",parent=styles["Heading1"],fontName=bold_name,fontSize=11,textColor=white,backColor=navy,borderPadding=5,spaceBefore=8,spaceAfter=5)
    story=[Paragraph("DiTuS Kablo Analizör™ - BOQ/BOM/RFQ Paketi",title),Paragraph(f"{escape(package.project_name)} · {escape(package.project_code)} · {escape(package.summary.status)}",body),Spacer(1,4*mm)]
    story.append(Paragraph("BOQ / BOM",banner))
    headers=("Kalem","Kategori","Tanım","Teknik özellik","Otomatik","Nihai","Birim","Durum","Dayanak")
    data=[[Paragraph(escape(h),table_header) for h in headers]]
    for item in package.lines:
        data.append([Paragraph(escape(str(x)),small) for x in (item.item_id,item.category,item.description,item.technical_specification,_fmt(item.auto_quantity),_fmt(item.final_quantity),item.unit,item.status,item.basis.formula)])
    widths=[18*mm,24*mm,38*mm,70*mm,18*mm,18*mm,13*mm,30*mm,45*mm]
    table=Table(data,colWidths=widths,repeatRows=1,splitByRow=True)
    table.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),navy),("TEXTCOLOR",(0,0),(-1,0),white),("FONTNAME",(0,0),(-1,0),bold_name),("GRID",(0,0),(-1,-1),0.3,line),("VALIGN",(0,0),(-1,-1),"TOP"),("ROWBACKGROUNDS",(0,1),(-1,-1),[white,colors.HexColor("#F8FAFB")]),("LEFTPADDING",(0,0),(-1,-1),2),("RIGHTPADDING",(0,0),(-1,-1),2),("TOPPADDING",(0,0),(-1,-1),2),("BOTTOMPADDING",(0,0),(-1,-1),2)])); story.append(table)
    story.append(PageBreak()); story.append(Paragraph("Makara Planı",banner))
    headers=("Makara","Kesim","Fire/pay","Yedek","Toplam","Azami","Bakiye","Aşım","Kesimler","Durum")
    data=[[Paragraph(escape(h),table_header) for h in headers]]
    for drum in package.drums:
        data.append([Paragraph(escape(str(x)),small) for x in (drum.drum_id,f"{drum.route_cut_length_m:.1f}",f"{drum.order_allowance_m:.1f}",f"{drum.spare_stock_length_m:.1f}",f"{drum.loaded_length_m:.1f}",f"{drum.maximum_length_m:.1f}",f"{drum.capacity_balance_m:.1f}",f"{drum.overload_m:.1f}",", ".join(f"{c.cut_id}:{c.required_cut_length_m:g}m" for c in drum.cuts),drum.assignment_status)])
    table=Table(data,colWidths=[18*mm,18*mm,20*mm,18*mm,18*mm,18*mm,20*mm,18*mm,75*mm,28*mm],repeatRows=1)
    table.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),navy),("TEXTCOLOR",(0,0),(-1,0),white),("GRID",(0,0),(-1,-1),0.3,line),("VALIGN",(0,0),(-1,-1),"TOP"),("ROWBACKGROUNDS",(0,1),(-1,-1),[white,colors.HexColor("#F8FAFB")])]))
    story.append(table); story.append(Paragraph("Varsayımlar ve Uyarılar",banner))
    for text in package.assumptions+package.warnings: story.append(Paragraph("• "+escape(text),body))
    path.parent.mkdir(parents=True,exist_ok=True); doc.build(story)


def write_procurement_package(
    package: ProcurementPackage,
    output_directory: str | Path,
    base_name: str | None = None,
    formats: Iterable[str] | None = None,
) -> dict[str, Path]:
    directory=Path(output_directory); directory.mkdir(parents=True,exist_ok=True)
    safe=re.sub(r"[^A-Za-z0-9_.-]+","_",base_name or f"{package.project_code}_BOQ_BOM_RFQ_v{__version__}").strip("_") or "ditus_procurement"
    requested=tuple(dict.fromkeys(str(x).lower().lstrip(".") for x in (formats or ("xlsx","csv","json","html","markdown","docx","pdf"))))
    result:dict[str,Path]={}
    for fmt in requested:
        if fmt=="json":
            p=directory/f"{safe}.json"; p.write_text(json.dumps(package.to_dict(),ensure_ascii=False,indent=2),encoding="utf-8"); result[fmt]=p
        elif fmt in {"md","markdown"}:
            p=directory/f"{safe}.md"; p.write_text(render_procurement_markdown(package),encoding="utf-8"); result["markdown"]=p
        elif fmt=="html":
            p=directory/f"{safe}.html"; p.write_text(render_procurement_html(package),encoding="utf-8"); result[fmt]=p
        elif fmt=="xlsx":
            p=directory/f"{safe}.xlsx"; _write_xlsx(package,p); result[fmt]=p
        elif fmt=="csv":
            csv_paths=_write_csv_bundle(package,directory,safe); result.update({f"csv_{k}":v for k,v in csv_paths.items()})
        elif fmt=="docx":
            p=directory/f"{safe}.docx"; _write_docx(package,p); result[fmt]=p
        elif fmt=="pdf":
            p=directory/f"{safe}.pdf"; _write_pdf(package,p); result[fmt]=p
        else:
            raise ProcurementInputError(f"Desteklenmeyen BOQ/BOM/RFQ formatı: {fmt}")
    return result
